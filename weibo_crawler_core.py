#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
微博PC端爬虫核心模块
====================
功能:
  1. 自动识别微博页面中的动态类名(CSS Modules 哈希类名,如 _body_m3n8j_63 -> _body_ecgcn_63)
  2. 收集指定博主在指定时间范围内的原创微博ID
  3. 逐条抓取详情并导出 Markdown 文件

类名识别策略(三级兜底):
  1. 优先使用配置文件 class_names.json 中保存的类名(上次成功使用的)
  2. 自动探测:通过 JS 在页面中按"语义特征"(时间链接、展开按钮、正文文本等)定位元素并提取类名
  3. 自动探测失败 -> 回调用户手动输入类名

用法(命令行):
  python weibo_crawler_cli.py --name 卢诗翰 --uid 3276099007 --start 2026-04-01 --end 2026-04-30

用法(GUI):
  python weibo_crawler_gui.py
"""

import os
import re
import sys
import json
import time
import logging
import random
import shutil
import urllib.parse
from datetime import datetime, timedelta

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import InvalidSelectorException
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.edge.service import Service as EdgeService
from webdriver_manager.microsoft import EdgeChromiumDriverManager

logger = logging.getLogger('weibo_crawler')

# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------

def ensure_logger(log_level=logging.INFO):
    """确保全局 logger 有输出 handler(避免重复添加)"""
    if not logger.handlers:
        handler = logging.StreamHandler(sys.stdout)
        handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
        logger.addHandler(handler)
    logger.setLevel(log_level)
    return logger


def _interruptible_sleep(seconds, stop_event=None):
    """分段睡眠,期间检查停止标志;支持优雅停止(每段 1 秒)"""
    if stop_event is None:
        time.sleep(seconds)
        return
    remaining = float(seconds)
    while remaining > 0:
        if stop_event.is_set():
            return
        step = min(1.0, remaining)
        time.sleep(step)
        remaining -= step


# ---------- 博主记录文件(DataPC/博主记录.txt) ----------

BLOGGER_RECORD_FILE = "博主记录.txt"


def blogger_record_path(data_root=None):
    """博主记录文件路径: 默认位于 DataPC 目录下"""
    root = data_root or os.path.join(app_dir(), "DataPC")
    return os.path.join(root, BLOGGER_RECORD_FILE)


def load_blogger_records(data_root=None):
    """读取博主记录文件,返回 {user_id: user_name}(文件不存在则返回空)

    文件格式: 每行 "微博ID 博主昵称"(空格分隔),支持 # 注释
    """
    records = {}
    path = blogger_record_path(data_root)
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith("#"):
                        continue
                    parts = line.split()
                    if len(parts) >= 2 and parts[0].isdigit():
                        records[parts[0]] = " ".join(parts[1:])
    except Exception as e:
        logger.warning(f"读取博主记录文件失败: {e}")
    return records


def save_blogger_record(user_id, user_name, data_root=None):
    """将博主写入记录文件(已存在则跳过),返回是否新增"""
    path = blogger_record_path(data_root)
    records = load_blogger_records(data_root)
    if user_id in records:
        return False
    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        # 修复: 若文件已存在但末尾没有换行(如手动编辑过最后一行),
        # 追加时会把新记录拼到上一行,导致记录解析错误。
        need_newline = False
        if os.path.exists(path):
            with open(path, "rb") as f:
                data = f.read()
            if data and not data.endswith((b"\n", b"\r")):
                need_newline = True
        with open(path, "a", encoding="utf-8") as f:
            if need_newline:
                f.write("\n")
            f.write(f"{user_id} {user_name}\n")
        logger.info(f"已记录博主: {user_name}({user_id}) -> {path}")
        return True
    except Exception as e:
        logger.warning(f"写入博主记录失败: {e}")
        return False


def app_dir():
    """获取应用程序所在目录

    - 源码运行: 脚本所在目录
    - PyInstaller 打包的 exe: exe 所在目录(而非临时解压目录 _MEIPASS)
    """
    if getattr(sys, "frozen", False):  # PyInstaller 打包环境
        return os.path.dirname(os.path.abspath(sys.executable))
    return os.path.dirname(os.path.abspath(__file__))


def resource_path(*parts):
    """获取资源文件路径

    - 源码运行: 脚本所在目录
    - PyInstaller 打包: 优先从临时解压目录 _MEIPASS 读取内置资源(只读);
      需要持久化的文件(如 class_names.json 保存)仍走 app_dir(exe 旁)
    """
    if getattr(sys, "frozen", False):
        meipass = getattr(sys, "_MEIPASS", "")
        if meipass:
            bundled = os.path.join(meipass, *parts)
            if os.path.exists(bundled):
                return bundled
    return os.path.join(app_dir(), *parts)


# ---------------------------------------------------------------------------
# 类名管理:配置文件 + 自动探测 + 手动输入兜底
# ---------------------------------------------------------------------------

class ClassNameManager:
    """微博页面类名管理

    管理的类名(键名):
      card       微博卡片容器  (div._body_xxx_63)
      time       发布时间链接  (a._time_xxx_33)
      name       博主用户名    (div._name_xxx_291)
      content    微博正文      (div._wbtext_xxx_14)
      stats_wrap 转评赞容器    (div._wrap_xxx_137)
      stats_num  转评赞数字    (span._num_xxx_46)

    配置保存在 class_names.json,格式:
      {"card": "_body_ecgcn_63", "time": "...", ...}
    """

    CONFIG_FILE = "class_names.json"

    # 默认类名:用户已知的最新值,配置缺失时作为初始尝试
    DEFAULT_CLASSES = {
        "card": "_body_ecgcn_63",        # 2026-08 微博卡片类名(用户提供)
        "time": "_time_1tpft_33",
        "name": "_name_1yc79_291",
        "content": "_wbtext_q1l14_14",
        "stats_wrap": "_wrap_198pe_137",
        "stats_num": "_num_198pe_46",
    }

    # 各键对应的元素标签,用于 CSS 选择器与验证
    TAG_MAP = {
        "card": "div",
        "time": "a",
        "name": "div",
        "content": "div",
        "stats_wrap": "div",
        "stats_num": "span",
    }

    def __init__(self, config_path=None, manual_callback=None, logger=None,
                 detected_callback=None):
        """
        manual_callback: 手动输入类名的回调函数
            签名: callback(key, current_value) -> str
            返回用户输入的类名(不含前缀),返回空串/None 表示用户放弃输入
        detected_callback: 类名更新(探测成功/手动输入生效)通知回调
            签名: callback(key, new_value) -> str
            用于把新类名同步到界面输入框/设置(修复 gui_settings 残留旧值)
        """
        self.logger = logger or logging.getLogger('weibo_crawler.classes')
        self.config_path = config_path or os.path.join(app_dir(), self.CONFIG_FILE)
        self.manual_callback = manual_callback
        self.detected_callback = detected_callback
        self.classes = dict(self.DEFAULT_CLASSES)
        self.load()

    def _notify_detected(self, key, value):
        """类名确定后通知外部(如 GUI 同步输入框与 gui_settings)"""
        if self.detected_callback:
            try:
                self.detected_callback(key, value)
            except Exception as e:
                self.logger.warning(f"类名更新通知失败: {e}")

    # ---------- 配置读写 ----------

    def load(self):
        """从配置文件加载类名,文件不存在或损坏时使用默认值"""
        try:
            if os.path.exists(self.config_path):
                with open(self.config_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    for key in self.DEFAULT_CLASSES:
                        val = data.get(key)
                        if val and isinstance(val, str):
                            self.classes[key] = val
                    self.logger.info(f"已从配置文件加载类名: {self.config_path}")
        except Exception as e:
            self.logger.warning(f"读取类名配置文件失败,使用默认类名: {e}")

    def save(self):
        """保存当前类名到配置文件"""
        try:
            with open(self.config_path, "w", encoding="utf-8") as f:
                json.dump(self.classes, f, ensure_ascii=False, indent=2)
            self.logger.info(f"类名已保存到: {self.config_path}")
            return True
        except Exception as e:
            self.logger.error(f"保存类名配置文件失败: {e}")
            return False

    def get(self, key):
        return self.classes.get(key, "")

    def set(self, key, value):
        value = (value or "").strip().lstrip(".")
        if value:
            self.classes[key] = value

    # ---------- 类名解析(核心):配置 -> 探测 -> 手动 ----------

    def resolve(self, driver, key, uid=None, username=None, no_data_check=None):
        """确保 key 对应的类名在当前页面有效

        流程:
          1. 当前配置的类名若能在页面找到有效元素 -> 直接使用
          2. 否则自动探测(JS 按语义特征查找)
          3. 探测失败 -> 若 no_data_check 回调返回 True(页面没有目标数据,
             如博主该时间段无微博),跳过手动输入直接返回 None
          4. 否则 -> 调用 manual_callback 让用户输入

        返回有效类名(不含点号),全部失败返回 None
        """
        tag = self.TAG_MAP.get(key, "div")
        current = self.get(key)

        # 第1步:验证现有类名
        if current:
            try:
                els = driver.find_elements(By.CSS_SELECTOR, f"{tag}.{current}")
                if els:
                    # 对 card / time 做附加验证(必须包含目标用户的微博链接)
                    if key == "card" and uid:
                        ok = any(self._element_has_user_link(el, uid) for el in els[:5])
                        if ok:
                            return current
                    elif key == "time" and uid:
                        ok = any(self._element_has_user_link(el, uid) for el in els[:5])
                        if ok:
                            return current
                    else:
                        return current
                self.logger.info(f"类名 {key}={current} 在页面中未找到有效元素,尝试自动探测")
            except InvalidSelectorException:
                # 修复: 非法选择器(如曾手动输入过 "1")不打印完整堆栈,一句话提示即可
                self.logger.warning(f"类名 {key}={current} 不是合法的CSS选择器,尝试自动探测")
            except Exception as e:
                self.logger.warning(f"验证类名 {key}={current} 时出错: {e}")

        # 第2步:自动探测
        detected = self._probe(driver, key, uid, username)
        if detected:
            self.set(key, detected)
            self.save()
            self.logger.info(f"已自动探测到类名 {key}: {detected}")
            self._notify_detected(key, detected)
            return detected

        # 第3步:若页面根本没有目标数据(如博主该时间段无微博),跳过手动输入
        if no_data_check is not None:
            try:
                if no_data_check():
                    self.logger.info(
                        f"页面中未发现该博主的微博数据,跳过类名 {key} 的手动输入")
                    return None
            except Exception as e:
                self.logger.warning(f"检查页面数据状态时出错: {e}")

        # 第4步:手动输入
        self.logger.warning(f"自动探测类名 {key} 失败,需要手动输入")
        if self.manual_callback:
            while True:
                try:
                    val = self.manual_callback(key, current or "")
                except KeyboardInterrupt:
                    return None
                val = (val or "").strip().lstrip(".")
                if not val:
                    self.logger.warning(f"用户放弃输入类名 {key}")
                    return None
                # 验证用户输入的类名
                try:
                    els = driver.find_elements(By.CSS_SELECTOR, f"{tag}.{val}")
                    if els:
                        if key in ("card", "time") and uid:
                            if any(self._element_has_user_link(el, uid) for el in els[:5]):
                                self.set(key, val)
                                self.save()
                                self._notify_detected(key, val)
                                return val
                            self.logger.warning(f"输入的类名 {val} 未找到目标用户的链接,请确认")
                            continue
                        self.set(key, val)
                        self.save()
                        self._notify_detected(key, val)
                        return val
                    self.logger.warning(f"输入的类名 {val} 在页面上找不到元素,请重新输入")
                except InvalidSelectorException:
                    # 修复: 非法选择器一句话提示,不打印完整堆栈
                    self.logger.warning(f"输入的类名 {val} 不是合法的CSS选择器,请重新输入")
                    continue
                except Exception as e:
                    self.logger.warning(f"验证手动输入的类名 {val} 时出错: {e}")
                    continue
        return None

    @staticmethod
    def _element_has_user_link(el, uid):
        """判断元素内部是否包含指向 weibo.com/{uid}/{mid} 的时间链接"""
        try:
            links = el.find_elements(By.CSS_SELECTOR, "a[href*='/{}']".format(uid))
            for link in links:
                href = link.get_attribute("href") or ""
                title = link.get_attribute("title") or ""
                if f"/{uid}/" in href and re.match(r"\d{4}-\d{1,2}-\d{1,2}", title):
                    return True
        except Exception:
            pass
        return False

    # ---------- 自动探测(JS) ----------

    def _probe(self, driver, key, uid=None, username=None):
        """按语义特征在页面中探测类名,返回类名或 None"""
        try:
            if key == "card":
                return driver.execute_script(self._JS_CARD, uid or "")
            if key == "time":
                return driver.execute_script(self._JS_TIME, uid or "")
            if key == "name":
                return driver.execute_script(self._JS_NAME, username or "")
            if key == "content":
                return driver.execute_script(self._JS_CONTENT)
            if key == "stats_wrap":
                return driver.execute_script(self._JS_STATS)["wrap"]
            if key == "stats_num":
                return driver.execute_script(self._JS_STATS)["num"]
        except Exception as e:
            self.logger.warning(f"自动探测类名 {key} 时出错: {e}")
        return None

    # 探测卡片类名:找指向 uid 的时间链接(a[title] 为日期),向上找 class 含 _body_ 的元素
    _JS_CARD = r"""
    return (function(uid){
      if(!uid) return null;
      var links = document.querySelectorAll('a');
      var counts = {};
      var bodyRe = /_body_[A-Za-z0-9]+_\d+/;
      for (var i=0;i<links.length;i++){
        var a = links[i];
        var href = a.getAttribute('href')||'';
        var title = a.getAttribute('title')||'';
        if (href.indexOf('/'+uid+'/')===-1) continue;
        if (!/^\d{4}-\d{1,2}-\d{1,2}/.test(title)) continue;
        var el = a;
        while (el && el!==document.body){
          var cls = el.getAttribute('class')||'';
          if (cls && typeof cls==='string'){
            var m = cls.match(bodyRe);
            if (m){ counts[m[0]]=(counts[m[0]]||0)+1; break; }
          }
          el = el.parentElement;
        }
      }
      var best=null,bn=0;
      for (var k in counts){ if(counts[k]>bn){bn=counts[k];best=k;} }
      return best;
    })(arguments[0])
    """

    # 探测时间链接类名:找指向 uid 且 title 为日期的 a,取 class 中 _time_ 开头者
    _JS_TIME = r"""
    return (function(uid){
      if(!uid) return null;
      var links = document.querySelectorAll('a');
      var counts = {};
      var timeRe = /_time_[A-Za-z0-9]+_\d+/;
      for (var i=0;i<links.length;i++){
        var a = links[i];
        var href = a.getAttribute('href')||'';
        var title = a.getAttribute('title')||'';
        if (href.indexOf('/'+uid+'/')===-1) continue;
        if (!/^\d{4}-\d{1,2}-\d{1,2}/.test(title)) continue;
        var cls = a.getAttribute('class')||'';
        if (cls && typeof cls==='string'){
          var m = cls.match(timeRe);
          if (m){ counts[m[0]]=(counts[m[0]]||0)+1; }
        }
      }
      var best=null,bn=0;
      for (var k in counts){ if(counts[k]>bn){bn=counts[k];best=k;} }
      return best;
    })(arguments[0])
    """

    # 探测用户名类名:找文本等于博主名的元素,class 中 _name_ 开头者
    _JS_NAME = r"""
    return (function(username){
      if(!username) return null;
      var els = document.querySelectorAll('div, span, h1, h2, h3');
      var counts = {};
      var nameRe = /_name_[A-Za-z0-9]+_\d+/;
      for (var i=0;i<els.length;i++){
        var el = els[i];
        var txt = (el.innerText||'').trim();
        if (txt !== username && txt.indexOf(username) === -1) continue;
        if (txt.length > username.length + 20) continue; // 避免命中大块文本容器
        var cls = el.getAttribute('class')||'';
        if (cls && typeof cls==='string'){
          var m = cls.match(nameRe);
          if (m){ counts[m[0]]=(counts[m[0]]||0)+1; }
        }
      }
      var best=null,bn=0;
      for (var k in counts){ if(counts[k]>bn){bn=counts[k];best=k;} }
      return best;
    })(arguments[0])
    """

    # 探测正文类名(详情页):class 含 _wbtext_ 且文本较长的元素
    _JS_CONTENT = r"""
    return (function(){
      var els = document.querySelectorAll('div');
      var counts = {};
      var re = /_wbtext_[A-Za-z0-9]+_\d+/;
      for (var i=0;i<els.length;i++){
        var el = els[i];
        var cls = el.getAttribute('class')||'';
        if (!cls || typeof cls!=='string') continue;
        var m = cls.match(re);
        if (!m) continue;
        var txt = (el.innerText||'').trim();
        if (txt.length >= 5){ counts[m[0]]=(counts[m[0]]||0)+1; }
      }
      var best=null,bn=0;
      for (var k in counts){ if(counts[k]>bn){bn=counts[k];best=k;} }
      return best;
    })()
    """

    # 探测转评赞容器/数字类名(详情页):i.woo-font--* 图标的祖先 div 中 _wrap_ 开头者,
    # 以及其内部 class 含 _num_ 的 span
    _JS_STATS = r"""
    return (function(){
      var icons = document.querySelectorAll('i.woo-font--retweet, i.woo-font--comment, i.woo-font--like, i.woo-font--like2');
      var wrapCounts = {}, numCounts = {};
      var wrapRe = /_wrap_[A-Za-z0-9]+_\d+/;
      var numRe = /_num_[A-Za-z0-9]+_\d+/;
      var wraps = [];
      for (var i=0;i<icons.length;i++){
        var el = icons[i].parentElement;
        while (el && el!==document.body){
          var cls = el.getAttribute('class')||'';
          if (cls && typeof cls==='string'){
            var m = cls.match(wrapRe);
            if (m){ wrapCounts[m[0]]=(wrapCounts[m[0]]||0)+1; wraps.push(el); break; }
          }
          el = el.parentElement;
        }
      }
      for (var j=0;j<wraps.length;j++){
        var spans = wraps[j].querySelectorAll('span');
        for (var k=0;k<spans.length;k++){
          var scls = spans[k].getAttribute('class')||'';
          if (scls && typeof scls==='string'){
            var m2 = scls.match(numRe);
            if (m2){ numCounts[m2[0]]=(numCounts[m2[0]]||0)+1; }
          }
        }
      }
      function bestOf(o){ var b=null,n=0; for (var k in o){ if(o[k]>n){n=o[k];b=k;} } return b; }
      return {wrap: bestOf(wrapCounts), num: bestOf(numCounts)};
    })()
    """


# ---------------------------------------------------------------------------
# 微博 PC 端爬虫(收集 ID + 导出 Markdown)
# ---------------------------------------------------------------------------

class WeiboPCCrawler:
    """微博 PC 端爬虫

    职责:
      - 收集指定用户指定时间范围内的原创微博 ID -> 保存 txt
      - 从 ID 列表抓取详情 -> 导出 Markdown 文件
      - 所有页面类名通过 ClassNameManager 动态解析
    """

    def __init__(self, headless=False, user_data_dir=None,
                 class_manager=None, wait_callback=None, manual_callback=None,
                 detected_callback=None):
        self.user_data_dir = user_data_dir
        self.headless = headless
        # manual_callback: 手动输入类名的回调,签名 (key, current) -> str
        self.manual_callback = manual_callback
        self.classes = class_manager or ClassNameManager(
            manual_callback=manual_callback, detected_callback=detected_callback)
        # wait_callback: 等待用户操作(如手动登录后继续)的回调,签名 (message) -> None
        self.wait_callback = wait_callback
        self.driver = self.setup_driver(headless)
        self.wait = WebDriverWait(self.driver, 10)

    # ---------- 浏览器驱动 ----------

    def setup_driver(self, headless):
        """设置 Edge 浏览器选项

        驱动管理策略(按优先级):
          1. Selenium Manager(Selenium 4.6+ 内置,自动下载匹配版本,无需联网配置)
          2. webdriver-manager(旧方案,失败时忽略)
          3. 系统默认驱动
        启动前会自动清理上次异常退出遗留的锁文件(DevToolsActivePort / Singleton*)
        """
        edge_options = EdgeOptions()
        edge_options.use_chromium = True

        if headless:
            edge_options.add_argument('--headless')

        edge_options.add_argument('--disable-gpu')
        edge_options.add_argument('--no-sandbox')
        edge_options.add_argument('--disable-dev-shm-usage')
        edge_options.add_argument('--window-size=1920,1080')
        edge_options.add_argument(
            '--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36 Edg/120.0.0.0')

        if self.user_data_dir:
            edge_options.add_argument(f"--user-data-dir={self.user_data_dir}")
            logger.info(f"使用用户数据目录: {self.user_data_dir}")

        edge_options.add_argument("--disable-blink-features=AutomationControlled")
        edge_options.add_experimental_option("excludeSwitches", ["enable-automation"])
        edge_options.add_experimental_option('useAutomationExtension', False)

        self._cleanup_stale_locks()

        # 方案1: Selenium Manager(内置,自动管理驱动)
        try:
            driver = webdriver.Edge(options=edge_options)
            driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
            return driver
        except Exception as e:
            logger.warning(f"使用 Selenium Manager 创建 Edge 驱动失败: {e}")

        # 方案2: webdriver-manager
        try:
            service = EdgeService(EdgeChromiumDriverManager().install())
            driver = webdriver.Edge(service=service, options=edge_options)
            driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
            return driver
        except Exception as e:
            logger.error(f"创建 Edge 驱动失败: {e}")
            return self.setup_driver_fallback(headless)

    def _cleanup_stale_locks(self):
        """清理上次异常退出遗留的浏览器锁文件(DevToolsActivePort / Singleton*)"""
        if not self.user_data_dir:
            return
        try:
            for name in ("DevToolsActivePort", "DevToolsActivePort.lock",
                         "SingletonLock", "SingletonCookie", "SingletonSocket"):
                path = os.path.join(self.user_data_dir, name)
                if os.path.exists(path):
                    try:
                        os.remove(path)
                        logger.info(f"已清理遗留锁文件: {path}")
                    except Exception as e:
                        logger.warning(f"清理锁文件失败(可能仍被占用): {path} - {e}")
        except Exception as e:
            logger.warning(f"清理遗留锁文件时出错: {e}")

    def setup_driver_fallback(self, headless):
        """备用方法:使用系统默认的 Edge 驱动"""
        try:
            edge_options = EdgeOptions()
            edge_options.use_chromium = True

            if headless:
                edge_options.add_argument('--headless')

            edge_options.add_argument('--disable-gpu')
            edge_options.add_argument('--no-sandbox')
            edge_options.add_argument('--disable-dev-shm-usage')
            edge_options.add_argument('--window-size=1920,1080')
            edge_options.add_argument(
                '--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36 Edg/120.0.0.0')

            if self.user_data_dir:
                edge_options.add_argument(f"--user-data-dir={self.user_data_dir}")

            edge_options.add_argument("--disable-blink-features=AutomationControlled")
            edge_options.add_experimental_option("excludeSwitches", ["enable-automation"])
            edge_options.add_experimental_option('useAutomationExtension', False)

            driver = webdriver.Edge(options=edge_options)
            driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
            return driver
        except Exception as e:
            logger.error(f"备用方法也失败: {e}")
            raise Exception("无法创建 Edge 驱动,请确保已安装 Edge 浏览器")

    # ---------- 登录 ----------

    def check_login_status(self):
        """检查是否已登录微博"""
        try:
            self.driver.get("https://weibo.com")
            time.sleep(3)
            login_elements = self.driver.find_elements(By.XPATH, "//a[contains(text(), '登录')]")
            if login_elements:
                logger.info("检测到未登录状态")
                return False
            user_elements = self.driver.find_elements(By.XPATH, "//a[contains(@href, '/u/')]")
            if user_elements:
                logger.info("检测到已登录状态")
                return True
            return False
        except Exception as e:
            logger.error(f"检查登录状态时出错: {e}")
            return False

    def manual_login(self):
        """手动登录微博,等待用户在浏览器中完成登录

        若当前为无头模式(headless),自动重启为有头模式,
        否则用户看不到浏览器窗口无法扫码登录。
        登录成功后若原本是无头模式,自动恢复为无头模式继续(不打扰用户)。
        """
        logger.info("请手动登录微博...")

        # 无头模式下无法显示登录界面,自动切换为有头模式
        was_headless = bool(self.headless)
        if was_headless:
            logger.warning("当前为无头模式,无法显示浏览器供登录,自动重启为有头模式...")
            try:
                self.driver.quit()
            except Exception:
                pass
            self.driver = self.setup_driver(headless=False)
            self.wait = WebDriverWait(self.driver, 10)
            self.headless = False
            logger.info("已重启为有头模式,请在弹出的浏览器窗口中完成登录")

        self.driver.get("https://weibo.com/login.php")
        if self.wait_callback:
            self.wait_callback("请在浏览器中完成登录,然后点击「确定」继续...")
        else:
            input("请在浏览器中完成登录,然后按回车键继续...")
        if self.check_login_status():
            logger.info("登录成功")
            # 原本是无头模式: 登录完成后恢复无头模式,继续自动爬取
            if was_headless:
                logger.info("登录完成,恢复无头模式继续...")
                try:
                    self.driver.quit()
                except Exception:
                    pass
                self.driver = self.setup_driver(headless=True)
                self.wait = WebDriverWait(self.driver, 10)
                self.headless = True
            return True
        logger.warning("登录可能未成功,请检查")
        return False

    # ---------- 博主校验 ----------

    # 微博用户不存在时的页面提示(文本片段,取页面正文匹配)
    USER_NOT_FOUND_MARKERS = ("用户不存在", "用户不能存在", "该用户不存在",
                              "此用户不存在", "无此用户", "账号不存在")

    def verify_user(self, user_id, user_name=""):
        """爬取前访问用户主页,校验:
        1) 用户(微博ID)是否存在 - 不存在时抛 RuntimeError(友好提示)
        2) 主页昵称是否与输入昵称匹配 - 不匹配时抛 RuntimeError(防止爬错人)

        返回主页解析到的昵称(可能为空字符串,表示无法解析,不做强校验)。
        """
        url = f"https://weibo.com/u/{user_id}"
        logger.info(f"正在校验博主是否存在: {url}")
        self.driver.get(url)
        time.sleep(4)
        body_text = ""
        try:
            body_text = (self.driver.find_element(By.TAG_NAME, "body").text or "")
        except Exception as e:
            logger.warning(f"读取主页正文失败: {e}")
        for marker in self.USER_NOT_FOUND_MARKERS:
            if marker in body_text:
                raise RuntimeError(
                    f"博主不存在: 微博ID {user_id} 对应的用户不存在(页面提示“{marker}”),"
                    f"请检查ID是否正确。可访问 {url} 确认。")
        # 昵称匹配: 主页标题形如 “@<昵称> 的个人主页”(实测格式,部分场景为“的微博主页”);
        # 标题为异步加载(初始为 Sina Visitor System / -微博),轮询等待其变为有效值(最多 ~8 秒)
        title = ""
        for _ in range(8):
            try:
                title = (self.driver.title or "").strip()
            except Exception:
                title = ""
            if re.search(r"(?:的个人主页|的微博主页|的微博)$", title):
                break
            time.sleep(1)
        page_name = ""
        try:
            m = re.search(r"^@?\s*(.+?)\s*(?:的个人主页|的微博主页|的微博)$", title)
            if m:
                page_name = m.group(1).strip()
        except Exception:
            pass
        # “我的”是访问自己主页时的标题,无法据此核对昵称,跳过
        if page_name and page_name == "我的":
            page_name = ""
        if page_name and user_name and user_name != user_id and page_name != user_name:
            raise RuntimeError(
                f"博主昵称与微博ID不匹配: 输入的昵称是“{user_name}”,"
                f"但该ID主页显示的昵称是“{page_name}”,请确认后重试。\n"
                f"可访问 {url} 查看实际情况。")
        return page_name

    # ---------- URL 构建 ----------

    @staticmethod
    def date_to_timestamp(date_str):
        """将日期字符串转换为微博时间戳(北京时间 0 点对应的 UTC 时间戳)"""
        try:
            dt = datetime.strptime(date_str, "%Y-%m-%d")
            utc_dt = dt - timedelta(hours=8)
            return int(utc_dt.timestamp())
        except Exception as e:
            logger.error(f"日期转换失败: {date_str}, 错误: {e}")
            return None

    def build_search_url(self, user_id, is_ori=1, is_forward=1, is_text=1, is_pic=1,
                         is_video=1, is_music=1, keyword="", start_date=None, end_date=None):
        """构建带高级搜索参数的微博用户页 URL"""
        base_url = f"https://weibo.com/u/{user_id}"
        params = {}
        if is_ori:
            params["is_ori"] = 1
        if is_forward:
            params["is_forward"] = 1
        if is_text:
            params["is_text"] = 1
        if is_pic:
            params["is_pic"] = 1
        if is_video:
            params["is_video"] = 1
        if is_music:
            params["is_music"] = 1
        if keyword:
            params["key_word"] = keyword
        if start_date:
            ts = self.date_to_timestamp(start_date)
            if ts:
                params["start_time"] = ts
        if end_date:
            ts = self.date_to_timestamp(end_date)
            if ts:
                params["end_time"] = ts
        if params:
            return f"{base_url}?{urllib.parse.urlencode(params)}"
        return base_url

    # ---------- 类名解析辅助 ----------

    def _page_has_user_links(self, uid):
        """诊断辅助:页面上是否存在指向该用户的微博时间链接

        用于区分两种情况:
          - 有链接但探测不到类名 -> 类名结构确实变了,需要手动输入
          - 没有链接 -> 该时间段内博主可能没有发表微博
        """
        try:
            return bool(self.driver.execute_script(self._JS_HAS_USER_LINKS, uid))
        except Exception as e:
            logger.warning(f"检查页面微博链接时出错: {e}")
            return False

    _JS_HAS_USER_LINKS = r"""
    return (function(uid){
      if(!uid) return false;
      var links = document.querySelectorAll('a');
      for (var i=0;i<links.length;i++){
        var a = links[i];
        var href = a.getAttribute('href')||'';
        var title = a.getAttribute('title')||'';
        if (href.indexOf('/'+uid+'/')===-1) continue;
        if (/^\d{4}-\d{1,2}-\d{1,2}/.test(title)) return true;
      }
      return false;
    })(arguments[0])
    """

    def resolve_card_class(self, uid):
        """解析微博卡片类名

        失败时先诊断页面状态,给出针对性提示:
          - 页面没有该博主的微博链接 -> 可能该时间段没有发表微博
          - 有链接但探测失败 -> 类名结构变化,需要手动输入
        """
        cls = self.classes.resolve(
            self.driver, "card", uid=uid,
            no_data_check=lambda: not self._page_has_user_links(uid))
        if not cls:
            if not self._page_has_user_links(uid):
                raise RuntimeError(
                    "页面上未找到该博主的微博:博主可能在此时间段内没有发表微博"
                    "(或页面尚未加载完成),请确认时间范围是否正确")
            raise RuntimeError(
                "无法确定微博卡片类名:自动探测失败且未获得手动输入,"
                "请检查页面类名是否变化后重试")
        return cls

    def resolve_time_class(self, uid):
        """解析时间链接类名,失败时同样先诊断页面状态"""
        cls = self.classes.resolve(
            self.driver, "time", uid=uid,
            no_data_check=lambda: not self._page_has_user_links(uid))
        if not cls:
            if not self._page_has_user_links(uid):
                raise RuntimeError(
                    "页面上未找到该博主的微博:博主可能在此时间段内没有发表微博"
                    "(或页面尚未加载完成),请确认时间范围是否正确")
            raise RuntimeError(
                "无法确定时间链接类名:自动探测失败且未获得手动输入,"
                "请检查页面类名是否变化后重试")
        return cls

    # ---------- 收集阶段 ----------

    def get_username(self, user_id, fallback_name=""):
        """从页面获取用户名"""
        try:
            time.sleep(2)
            name_cls = self.classes.resolve(self.driver, "name", username=fallback_name)
            if name_cls:
                els = self.driver.find_elements(By.CSS_SELECTOR, f"div.{name_cls}")
                for el in els:
                    txt = el.text.strip()
                    if txt:
                        logger.info(f"获取到用户名: {txt}")
                        return txt
            return fallback_name or f"user_{user_id}"
        except Exception as e:
            logger.error(f"获取用户名失败: {e}")
            return fallback_name or f"user_{user_id}"

    def click_search_button(self, wait_after=5):
        """点击页面上的搜索按钮,确保高级搜索条件(时间/类型)生效

        微博页面首次加载时 URL 参数可能未生效,点击"搜索"按钮后
        页面才会只展示符合条件(如时间范围内)的微博。
        支持多种按钮结构,并容忍文本两侧空白。
        """
        selectors = [
            # 用户提供的结构: <span class="woo-button-content"> 搜索 </span>
            "//button[.//span[normalize-space(text())='搜索']]",
            "//button[.//span[contains(@class, 'woo-button-content') and "
            "normalize-space(.)='搜索']]",
            # 兜底: 任意含"搜索"文本的按钮
            "//button[.//*[normalize-space(text())='搜索']]",
            "//*[@role='button' and .//*[normalize-space(text())='搜索']]",
        ]
        for selector in selectors:
            try:
                buttons = self.driver.find_elements(By.XPATH, selector)
                if buttons:
                    buttons[0].click()
                    logger.info("已点击搜索按钮,等待搜索结果刷新...")
                    time.sleep(wait_after)
                    return True
            except Exception as e:
                logger.warning(f"尝试选择器 {selector[:50]} 时出错: {e}")
        logger.warning("找不到搜索按钮,直接使用 URL 参数(时间过滤可能未生效)")
        return False

    def collect_weibo_ids(self, user_id, start_date=None, end_date=None,
                          max_count=500, keyword="", is_ori=1, is_forward=0,
                          is_text=1, is_pic=1, is_video=1, is_music=1,
                          user_name=None, min_words=0, stop_event=None):
        """收集指定用户指定时间范围内的微博 ID(原创/转发可选)

        min_words>0 时在列表页做粗过滤: 卡片预览字数低于该值的微博
        不收集(不进详情页);精确过滤仍由导出阶段(min_words)负责。

        返回 (username, weibo_ids)
        """
        # 登录检查
        if not self.check_login_status():
            logger.warning("未检测到登录状态,尝试手动登录")
            if not self.manual_login():
                logger.error("登录失败,无法继续")
                return None, []

        # 爬取前校验博主: ID 是否存在、昵称是否与 ID 匹配(避免爬错人/误提示"该时间段无微博")
        try:
            self.verify_user(user_id, user_name or "")
        except RuntimeError:
            raise  # 留给 run_task 统一处理
        except Exception as e:
            logger.warning(f"博主校验未完成(继续爬取): {e}")

        search_url = self.build_search_url(
            user_id, is_ori, is_forward, is_text, is_pic,
            is_video, is_music, keyword, start_date, end_date)
        logger.info(f"正在访问URL: {search_url}")
        self.driver.get(search_url)
        time.sleep(5)

        # 关键: 点击页面上的"搜索"按钮,让时间/类型过滤条件真正生效。
        # 微博页面首次加载时 URL 参数可能未应用,不点击会抓到范围外的微博。
        self.click_search_button()

        # 解析类名(在列表页上自动探测);若失败,先刷新页面再重试一次
        # (避免懒加载导致页面暂无卡片而误判"该时间段没有微博")
        try:
            card_cls = self.resolve_card_class(user_id)
        except RuntimeError as e:
            logger.warning(f"首次解析卡片类名失败: {e};刷新页面重试一次...")
            self.driver.refresh()
            time.sleep(5)
            card_cls = self.resolve_card_class(user_id)

        try:
            time_cls = self.resolve_time_class(user_id)
        except RuntimeError as e:
            logger.warning(f"首次解析时间类名失败: {e};刷新页面重试一次...")
            self.driver.refresh()
            time.sleep(5)
            time_cls = self.resolve_time_class(user_id)

        logger.info(f"使用的类名 -> 卡片: {card_cls}, 时间链接: {time_cls}")

        # 获取用户名(传入预期用户名以便探测)
        username = self.get_username(user_id, fallback_name=user_name or user_id)

        # 边滚动边收集微博 ID
        weibo_ids = self._scroll_and_collect(user_id, card_cls, time_cls, max_count,
                                             min_words=min_words, stop_event=stop_event)
        return username, weibo_ids

    def _scroll_and_collect(self, user_id, card_cls, time_cls, max_count=500,
                            min_words=0, stop_event=None):
        """渐进式滚动收集微博 ID"""
        weibo_ids = []
        processed_ids = set()
        scroll_attempts = 0
        max_scroll_attempts = 200
        no_progress_count = 0
        max_no_progress = 5

        current_scroll_position = 0
        last_scroll_height = self.driver.execute_script("return document.body.scrollHeight")
        card_selector = f"div.{card_cls}"

        while (len(weibo_ids) < max_count and
               scroll_attempts < max_scroll_attempts and
               no_progress_count < max_no_progress):

            if stop_event is not None and stop_event.is_set():
                logger.info("收到停止请求,停止收集微博ID(已找到 %d 条)" % len(weibo_ids))
                break

            new_ids = self._extract_ids_from_cards(card_selector, time_cls, user_id,
                                                    processed_ids, min_words=min_words)
            new_count = 0
            for wid in new_ids:
                if wid not in weibo_ids:
                    weibo_ids.append(wid)
                    processed_ids.add(wid)
                    new_count += 1
                    logger.info(f"找到符合条件的微博ID: {wid} (总数: {len(weibo_ids)})")
                    if len(weibo_ids) >= max_count:
                        break

            if new_count == 0:
                no_progress_count += 1
            else:
                no_progress_count = 0

            if len(weibo_ids) >= max_count:
                break

            window_height = self.driver.execute_script("return window.innerHeight")
            scroll_amount = window_height * random.uniform(0.4, 0.5)
            new_scroll_position = min(current_scroll_position + scroll_amount, last_scroll_height)
            self.driver.execute_script(f"window.scrollTo(0, {new_scroll_position});")
            time.sleep(random.uniform(1, 1.5))

            current_scroll_position = new_scroll_position
            new_scroll_height = self.driver.execute_script("return document.body.scrollHeight")
            if new_scroll_height > last_scroll_height:
                last_scroll_height = new_scroll_height
                no_progress_count = 0

            scroll_attempts += 1
            if current_scroll_position >= last_scroll_height - window_height:
                logger.info("已到达或接近页面底部")
                no_progress_count = max_no_progress

        logger.info(f"滚动完成,共尝试 {scroll_attempts} 次,找到 {len(weibo_ids)} 个微博ID")
        return weibo_ids[:max_count]

    def _extract_ids_from_cards(self, card_selector, time_cls, user_id, processed_ids,
                                min_words=0):
        """从当前页面卡片中提取微博 ID(含展开按钮、非转发)

        min_words>0 时做列表页粗过滤: 卡片正文预览字数低于 min_words
        的微博不收集(避免进入详情页,节省时间;最终以详情页精确过滤为准)
        """
        weibo_ids = []
        try:
            cards = self.driver.find_elements(By.CSS_SELECTOR, card_selector)
            logger.info(f"当前可见 {len(cards)} 个微博卡片")
        except Exception as e:
            logger.warning(f"查找微博卡片失败: {e}")
            return weibo_ids

        time_selector = f"a.{time_cls}"
        content_cls = self.classes.get("content")  # 列表页正文预览类名(可能为空)
        for i, card in enumerate(cards):
            try:
                # 必须有"展开"按钮(说明是长文)
                expand_buttons = card.find_elements(
                    By.XPATH, ".//span[contains(@class, 'expand') and contains(text(), '展开')]")
                if not expand_buttons:
                    continue
                # 跳过转发微博
                repost_elements = card.find_elements(
                    By.XPATH, ".//div[contains(@class, 'repost') or contains(text(), '转发微博')]")
                if repost_elements:
                    continue
                # 列表页粗过滤: 预览正文字数低于 min_words 的不收集
                if min_words > 0 and content_cls:
                    preview_els = card.find_elements(By.CSS_SELECTOR, f"div.{content_cls}")
                    preview = (preview_els[0].text or "") if preview_els else ""
                    if len(preview) < min_words:
                        logger.info(f"列表预览字数 {len(preview)} < {min_words},"
                                    f"粗过滤跳过该卡片")
                        continue
                # 时间链接
                time_links = card.find_elements(By.CSS_SELECTOR, time_selector)
                if not time_links:
                    continue
                href = time_links[0].get_attribute("href") or ""
                if f"/{user_id}/" in href:
                    parts = href.split(f"/{user_id}/")
                    if len(parts) > 1:
                        wid = parts[1].split("?")[0]
                        if wid and wid not in processed_ids:
                            weibo_ids.append(wid)
                            processed_ids.add(wid)
            except Exception as e:
                logger.warning(f"处理微博卡片 {i} 时出错: {e}")
                continue
        return weibo_ids

    # ---------- 详情阶段 ----------

    def get_weibo_detail(self, weibo_id, user_id):
        """获取单条微博的详细信息(打开新标签页)"""
        url = f"https://weibo.com/{user_id}/{weibo_id}"
        logger.info(f"正在访问微博: {url}")

        try:
            self.driver.execute_script(f"window.open('{url}');")
            self.driver.switch_to.window(self.driver.window_handles[-1])
            self.wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
            time.sleep(3)

            # 在详情页解析正文/转评赞类名(自动探测,失败则手动输入)
            content_cls = self.classes.resolve(self.driver, "content")
            self.classes.resolve(self.driver, "stats_wrap")
            self.classes.resolve(self.driver, "stats_num")

            publish_time = self.extract_publish_time()
            topics = self.extract_topics()
            content = self.extract_content(content_cls)
            repost_count, comment_count, like_count = self.extract_stats()
            images = self.extract_images()
            videos = self.extract_videos()
            save_time = datetime.now().strftime("%y-%m-%d %H:%M")

            return {
                "url": url,
                "publish_time": publish_time,
                "topics": topics,
                "content": content,
                "repost_count": repost_count,
                "comment_count": comment_count,
                "like_count": like_count,
                "images": images,
                "videos": videos,
                "save_time": save_time,
                "weibo_id": weibo_id,
            }
        except Exception as e:
            logger.error(f"获取微博内容时出现错误: {e}")
            return None
        finally:
            if len(self.driver.window_handles) > 1:
                self.driver.close()
                self.driver.switch_to.window(self.driver.window_handles[0])

    def extract_images(self):
        """提取微博正文配图的大图 URL(详情页)

        策略:
          1. 只在正文卡片(div.{card_cls},如 _body_ecgcn_63)范围内查找,
             避免抓到右侧用户卡片/推荐位等卡片外的图片
          2. 只提取配图类名:woo-picture-img(首图)与 _focusImg_*(其余图),
             过滤头像(woo-avatar-img)、VIP图标(woo-icon-vipimg)等
          3. 缩略图标记(orj360/mw690 等)替换为 large,取大图
        """
        images = []
        try:
            card_cls = self.classes.get("card")
            if not card_cls:
                return images
            cards = self.driver.find_elements(By.CSS_SELECTOR, f"div.{card_cls}")
            if not cards:
                return images
            # 详情页通常只有一个正文卡片;若有多个,取内容最长的(真正的正文)
            card = max(cards, key=lambda c: len(c.get_attribute("outerHTML") or ""))

            img_els = card.find_elements(By.CSS_SELECTOR, "img")
            seen = set()
            for img in img_els:
                src = img.get_attribute("src") or ""
                cls = img.get_attribute("class") or ""
                if not src or src.startswith("data:"):
                    continue
                # 只匹配配图类名;跳过头像/图标/视频占位图
                is_pic = ("woo-picture-img" in cls) or ("focusImg" in cls)
                if not is_pic:
                    continue
                url = self._clean_image_url(src)
                if url and url not in seen and len(seen) < 20:
                    seen.add(url)
                    images.append(url)
        except Exception as e:
            logger.warning(f"提取图片失败: {e}")
        return images

    @staticmethod
    def _clean_image_url(src):
        """清洗图片 URL:去掉查询参数/尺寸后缀,缩略图标记替换为 large 大图"""
        url = re.sub(r"\?.*$", "", src)          # 去掉 ?KID= 等查询参数
        url = re.sub(r"!\w+", "", url)           # 去掉 !thumb 等后缀
        # 缩略图标记 -> 大图: orj360/mw690/bmiddle/thumb150/square 等
        url = re.sub(r"/(?:orj360|mw690|bmiddle|thumb150|square)/",
                     "/large/", url)
        url = re.sub(r"#.*$", "", url)
        return url.strip()

    def extract_videos(self):
        """提取微博正文中的视频 URL(详情页)"""
        videos = []
        try:
            video_els = self.driver.find_elements(By.CSS_SELECTOR, "video source, video")
            seen = set()
            for v in video_els:
                src = v.get_attribute("src") or ""
                if not src:
                    continue
                if src and src not in seen and len(seen) < 10:
                    seen.add(src)
                    videos.append(src)
        except Exception as e:
            logger.warning(f"提取视频失败: {e}")
        return videos

    def extract_publish_time(self):
        """提取发布时间"""
        selectors = []
        time_cls = self.classes.get("time")
        if time_cls:
            selectors.append(f"a.{time_cls}")
        selectors += ["a[title*='-']", "span.time", "div.wb-info span"]
        for selector in selectors:
            try:
                els = self.driver.find_elements(By.CSS_SELECTOR, selector)
                if els:
                    txt = els[0].text.strip()
                    if txt:
                        return txt
                    title = els[0].get_attribute("title")
                    if title:
                        return title
            except Exception:
                continue
        return "未知时间"

    def extract_topics(self):
        """提取词条(话题)"""
        try:
            topics = []
            for el in self.driver.find_elements(By.CSS_SELECTOR, "a[href*='weibo?q=%23']"):
                txt = el.text.strip()
                if txt and txt.startswith('#') and txt.endswith('#'):
                    topics.append(txt)
            return " ".join(topics) if topics else "无"
        except Exception as e:
            logger.error(f"提取词条失败: {e}")
            return "无"

    def extract_content(self, content_cls=None):
        """提取正文内容"""
        selectors = []
        if content_cls:
            selectors.append(f"div.{content_cls}")
        selectors += ["div.wbtext", "div.weibo-text", "div.text"]
        for selector in selectors:
            try:
                els = self.driver.find_elements(By.CSS_SELECTOR, selector)
                if els:
                    txt = els[0].text.strip()
                    if txt:
                        return txt
            except Exception:
                continue
        # 兜底:整个页面文本
        try:
            body = self.driver.find_element(By.TAG_NAME, "body").text
            return body[:1000] if len(body) > 1000 else body
        except Exception as e:
            logger.error(f"提取微博内容时出现错误: {e}")
            return "无法提取内容"

    def extract_stats(self):
        """提取转发、评论、点赞数(多策略:动态类名优先,旧选择器兜底)"""
        repost_count = comment_count = like_count = "0"
        wrap_cls = self.classes.get("stats_wrap")
        num_cls = self.classes.get("stats_num")

        # 策略1: 通过图标 + 动态 wrap/num 类名提取转发/评论/点赞
        try:
            if wrap_cls and num_cls:
                icon_to_key = {
                    "woo-font--retweet": "repost",
                    "woo-font--comment": "comment",
                    "woo-font--like": "like",
                    "woo-font--like2": "like",
                }
                for icon_cls, key in icon_to_key.items():
                    try:
                        icons = self.driver.find_elements(By.CSS_SELECTOR, f"i.{icon_cls}")
                        if not icons:
                            continue
                        parent = icons[0].find_element(
                            By.XPATH, f"./ancestor::div[contains(@class, '{wrap_cls}')]")
                        nums = parent.find_elements(By.CSS_SELECTOR, f"span.{num_cls}")
                        if nums:
                            val = nums[0].text.strip()
                            if key == "repost":
                                repost_count = val
                            elif key == "comment":
                                comment_count = val
                            else:
                                like_count = val
                    except Exception:
                        continue
        except Exception as e:
            logger.warning(f"提取统计数据时出现异常: {e}")

        # 策略2(兜底): 使用旧的稳定选择器
        # 点赞数:button.woo-like-main span.woo-like-count
        if like_count == "0":
            try:
                like_els = self.driver.find_elements(
                    By.CSS_SELECTOR, "button.woo-like-main span.woo-like-count")
                if like_els:
                    val = like_els[0].text.strip()
                    if val:
                        like_count = val
            except Exception:
                pass
        # 转发/评论:通过 wrap/num 类名直接匹配
        if repost_count == "0" and wrap_cls and num_cls:
            try:
                wraps = self.driver.find_elements(By.CSS_SELECTOR, f"div.{wrap_cls}")
                for w in wraps[:6]:
                    try:
                        cls = w.get_attribute("class") or ""
                        num_els = w.find_elements(By.CSS_SELECTOR, f"span.{num_cls}")
                        if not num_els:
                            continue
                        val = num_els[0].text.strip()
                        if "retweet" in cls and repost_count == "0":
                            repost_count = val
                        elif "cur" in cls and comment_count == "0":
                            comment_count = val
                    except Exception:
                        continue
            except Exception:
                pass

        return repost_count, comment_count, like_count

    # ---------- Markdown 导出 ----------

    @staticmethod
    def generate_markdown(detail):
        """根据详情生成 Markdown 内容(URL 带超链接,便于阅读器直接点击)"""
        md = f">URL：[微博原文链接]({detail['url']})\n"
        md += f">发布时间：{detail['publish_time']}\n"
        md += f">词条：{detail['topics']}\n"
        md += f">正文字数：{len(detail['content'])}字符\n\n"
        md += "---\n\n正文：\n"
        md += f"{detail['content']}\n\n"
        md += "---\n\n"
        md += f">转发数：{detail['repost_count']}\n"
        md += f">评论数：{detail['comment_count']}\n"
        md += f">点赞数：{detail['like_count']}\n"
        md += f">保存时间：{detail['save_time']}"
        if detail.get("ai_quality") is not None:
            md += f"\n>高质量可信度：{detail['ai_quality']}%"
        return md

    def export_markdowns(self, weibo_ids, user_id, user_name, base_dir,
                         min_interval=3, max_interval=8, progress_callback=None,
                         overwrite=False, month_subdirs=True,
                         download_images=False, download_videos=False,
                         export_format="md", skip_existing=False, min_words=0,
                         ai_enabled=False, ai_classifier=None, ai_config=None,
                         ai_usage_callback=None, ai_rename=False, ai_root=None,
                         ai_copy_media=True, stop_event=None,
                         stats_store=None, force_rejudge=False):
        """逐条抓取详情并导出(支持 md / docx 格式)

        参数:
          min_interval / max_interval  每条微博之间的随机等待秒数范围
          overwrite                    同名文件直接覆盖(重跑场景)
          month_subdirs                按 YYYY-MM 月份子文件夹保存
          download_images/videos       是否下载图片/视频到本地
          export_format                "md" 或 "docx"
          skip_existing                已存在同ID文件时跳过(不重复抓取)
          min_words                    正文字数低于该值的文章不导出(0=不限制)
          ai_enabled                   启用 AI 实时判断
          ai_classifier                AIClassifier 实例(ai_enabled 时必传)
          ai_config                    AIConfig 实例(取高质量阈值等)
          ai_usage_callback            每次AI调用后回调(token增量), 用于显示用量
          ai_rename                    AI通过的文章同时生成标题并复制到 AI分类 目录
          ai_root                      AI分类输出根目录(默认 程序目录/AI分类)
          ai_copy_media                AI分类副本是否连带复制图片/视频
                                       (False 时md副本改为引用原网络链接)
          stats_store                  WeiboStatsStore 实例(状态记录,可选)
          force_rejudge                重跑时对已判定过的文章重新调用 AI(默认复用)

        返回 (成功数, 失败数, 跳过数, AI过滤数)
        """
        output_dir = base_dir  # 已由调用方创建
        ok_count = 0
        fail_count = 0
        skipped_count = 0
        ai_skipped_count = 0
        ai_quality_threshold = 80
        if ai_enabled and ai_config is not None:
            try:
                ai_quality_threshold = int(ai_config.get("quality_threshold", 80))
            except (TypeError, ValueError):
                ai_quality_threshold = 80
        total = len(weibo_ids)

        for i, weibo_id in enumerate(weibo_ids):
            # 修复: 支持优雅停止(在每两条之间检查标志,当前条目正常收尾)
            if stop_event is not None and stop_event.is_set():
                logger.info(f"收到停止请求,停止导出(已完成 {ok_count} 条)")
                break
            logger.info(f"正在处理第 {i + 1}/{total} 个微博: {weibo_id}")
            if progress_callback:
                progress_callback(i + 1, total, weibo_id)

            # 跳过已存在文件(按微博ID匹配,避免重复抓取)
            if skip_existing and self._weibo_file_exists(
                    output_dir, weibo_id, export_format):
                # 若勾选了下载图片/视频,检查媒体是否完整:
                # 文章存在但媒体缺失(如之前只爬文章没下图片)时,
                # 重新抓详情仅补下载媒体,不重写文章文件
                month_dir = self._find_weibo_dir(output_dir, weibo_id, export_format)
                need_media = False
                if download_images and month_dir and not self._media_exists(
                        month_dir, weibo_id, "images"):
                    need_media = True
                if download_videos and month_dir and not self._media_exists(
                        month_dir, weibo_id, "videos"):
                    need_media = True

                if need_media:
                    logger.info(f"文章已存在,检查媒体完整性: {weibo_id}")
                    detail = self.get_weibo_detail(weibo_id, user_id)
                    if detail:
                        if download_images:
                            if detail.get("images"):
                                self._download_media(
                                    detail["images"], month_dir, "images", weibo_id)
                            else:
                                # 确认无图,写标记避免下次重复检查
                                self._mark_no_media(month_dir, weibo_id, "images")
                                logger.info(f"确认无图片,已写标记: {weibo_id}")
                        if download_videos:
                            if detail.get("videos"):
                                self._download_media(
                                    detail["videos"], month_dir, "videos", weibo_id)
                            else:
                                self._mark_no_media(month_dir, weibo_id, "videos")
                                logger.info(f"确认无视频,已写标记: {weibo_id}")
                        # 更新状态记录中的媒体标记(供后续增量/更新转评赞使用)
                        if stats_store:
                            new_media = {}
                            if download_images:
                                new_media["images"] = bool(detail.get("images"))
                            if download_videos:
                                new_media["videos"] = bool(detail.get("videos"))
                            if new_media:
                                stats_store.set_record(
                                    weibo_id, media=new_media,
                                    publish=detail.get("publish_time", ""))
                                stats_store.save()
                        skipped_count += 1  # 文章未重写,仍计入跳过
                        continue
                    # 详情获取失败则继续正常跳过
                skipped_count += 1
                logger.info(f"已存在同ID文件,跳过: {weibo_id}")
                if stats_store:
                    stats_store.set_record(weibo_id, exported=True)
                    stats_store.save()
                continue

            # AI 判定复用(修复 AI 评分漂移): 之前被 AI 判定过且未导出的文章,
            # 默认直接复用上次分数,不再重复抓详情+调用大模型;force_rejudge 时重新判断
            rec = stats_store.get(weibo_id) if stats_store else None
            prev_ai = ((rec or {}).get("ai") or {})
            ai_quality = None
            ai_title = None
            if (ai_enabled and ai_classifier is not None and not force_rejudge
                    and prev_ai and prev_ai.get("quality") is not None
                    and not (rec or {}).get("exported")):
                prev_q = int(prev_ai["quality"])
                if prev_q < ai_quality_threshold:
                    ai_skipped_count += 1
                    logger.info(f"复用上次AI判定(可信度{prev_q}%<{ai_quality_threshold}%),"
                                f"跳过: {weibo_id}")
                    continue
                logger.info(f"复用上次AI判定(可信度{prev_q}%): {weibo_id}")
                ai_quality = prev_q

            detail = self.get_weibo_detail(weibo_id, user_id)
            if not detail:
                fail_count += 1
                logger.error(f"无法获取微博 {weibo_id} 的详情")
                continue

            # 最低字数过滤: 正文字数低于 min_words 时不导出
            if min_words > 0 and len(detail.get("content", "")) < min_words:
                skipped_count += 1
                logger.info(f"正文字数低于 {min_words},跳过: {weibo_id}")
                continue

            # AI 实时判断: 高质量可信度低于阈值则不导出(也不下载媒体)
            if ai_enabled and ai_classifier is not None and ai_quality is None:
                try:
                    _, _, quality_prob, usage = ai_classifier.classify(
                        detail.get("content", ""))
                    ai_quality = int(quality_prob or 0)
                    tokens = int(usage.get("total_tokens", 0) or 0)
                    if ai_usage_callback:
                        ai_usage_callback(tokens)
                    if ai_quality < ai_quality_threshold:
                        ai_skipped_count += 1
                        logger.info(
                            f"AI判定非高质量(高质量可信度{ai_quality}% < "
                            f"{ai_quality_threshold}%),跳过且不下载媒体: {weibo_id}")
                        if stats_store:
                            stats_store.set_record(weibo_id,
                                                   ai={"quality": ai_quality,
                                                       "checked_at": _now_str()})
                            stats_store.save()
                        continue
                    logger.info(f"AI判定高质量(高质量可信度{ai_quality}%): {weibo_id}")
                    # 可选: AI 总结生成标题(用于复制到 AI分类 目录时重命名)
                    if ai_rename:
                        try:
                            ai_title, usage2 = ai_classifier.summarize_title(
                                detail.get("content", ""))
                            tokens2 = int(usage2.get("total_tokens", 0) or 0)
                            if ai_usage_callback:
                                ai_usage_callback(tokens2)
                        except Exception as e:
                            logger.warning(f"AI总结失败(不影响导出): {weibo_id}: {e}")
                            ai_title = None
                except Exception as e:
                    # AI 失败时按通过处理,避免卡住整个爬取
                    logger.warning(f"AI实时判断失败({e}),按通过处理: {weibo_id}")
                    ai_quality = None

            # 文件名: 用户名_日期_微博ID.ext (日期来自发布时间)
            save_time = detail['publish_time']
            date_part = save_time.split()[0] if save_time and save_time != "未知时间" else "unknown"
            ext = ".docx" if export_format == "docx" else ".md"
            output_filename = f"{user_name}_{date_part}_{weibo_id}{ext}"

            # 按月保存: 根据发布时间归入 YYYY-MM 子文件夹
            if month_subdirs:
                month_dir = self._publish_month_dir(save_time, output_dir)
            else:
                month_dir = output_dir
            os.makedirs(month_dir, exist_ok=True)
            output_path = os.path.join(month_dir, output_filename)

            if not overwrite:
                counter = 1
                original = output_path
                while os.path.exists(output_path):
                    name, ext0 = os.path.splitext(original)
                    output_path = f"{name}_{counter}{ext0}"
                    counter += 1

            # 下载图片/视频到本地(可选,先下载以便 md 引用本地文件)
            # 修复: 确认无图/无视频时立即写 .nomedia 标记,
            # 否则重跑时会被当作"媒体缺失"再抓一次详情(浪费且日志误导)
            local_images = []
            if download_images and detail.get("images"):
                local_images = self._download_media(
                    detail["images"], month_dir, "images",
                    detail.get("weibo_id", ""))
            elif download_images:
                self._mark_no_media(month_dir, weibo_id, "images")
                logger.info(f"确认无图片,已写标记: {weibo_id}")
            local_videos = []
            if download_videos and detail.get("videos"):
                local_videos = self._download_media(
                    detail["videos"], month_dir, "videos",
                    detail.get("weibo_id", ""))
            elif download_videos:
                self._mark_no_media(month_dir, weibo_id, "videos")
                logger.info(f"确认无视频,已写标记: {weibo_id}")

            # 生成内容并写入
            try:
                if ai_quality is not None:
                    detail["ai_quality"] = ai_quality
                if export_format == "docx":
                    self._write_docx(detail, output_path, month_dir,
                                     local_images, local_videos,
                                     download_images, download_videos)
                else:
                    md_content = self.generate_markdown(detail)
                    if detail.get("images"):
                        md_content += "\n\n### 图片\n"
                        for idx, img_url in enumerate(detail["images"], 1):
                            if idx <= len(local_images) and local_images[idx - 1]:
                                md_content += f"\n![图片{idx}]({local_images[idx - 1]})\n"
                            else:
                                md_content += f"\n![图片{idx}]({img_url})\n"
                    if detail.get("videos"):
                        md_content += "\n\n### 视频\n"
                        for i2, v_url in enumerate(detail["videos"]):
                            if i2 < len(local_videos) and local_videos[i2]:
                                md_content += f"\n[视频]({local_videos[i2]})\n"
                            else:
                                md_content += f"\n[视频链接]({v_url})\n"
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(md_content)
            except Exception as e:
                logger.error(f"写入文件失败 {output_path}: {e}")
                fail_count += 1
                continue

            ok_count += 1
            logger.info(f"已保存: {os.path.basename(output_path)}")

            # AI 总结重命名: 复制一份到 AI分类/AI_<博主>_高质量/<年月>/ (保留原文件)
            if ai_title and ai_root:
                self._copy_to_ai_dir(output_path, user_name, save_time,
                                     ai_title, ai_root,
                                     copy_media=ai_copy_media, detail=detail)
                # 状态记录: 标记"已实时AI分类"(事后AI分类据此跳过,避免内容相同的重复文件)
                if stats_store:
                    stats_store.set_record(weibo_id, ai_copied=True,
                                           ai_title=ai_title, ai_copied_at=_now_str())
                    stats_store.save()

            # 更新状态记录: 导出结果/AI判定/转评赞/媒体(V0.5.2 状态记录)
            if stats_store:
                stats_store.set_record(
                    weibo_id,
                    exported=True,
                    file=os.path.basename(output_path),
                    publish=save_time,
                    ai={"quality": ai_quality, "checked_at": _now_str()}
                    if ai_quality is not None else None,
                    stats={"repost": detail["repost_count"],
                           "comment": detail["comment_count"],
                           "like": detail["like_count"],
                           "checked_at": _now_str()},
                    media={"images": bool(local_images),
                           "videos": bool(local_videos)})
                stats_store.save()

            if i < total - 1:
                sleep_time = random.randint(min_interval, max_interval)
                logger.info(f"等待 {sleep_time} 秒后处理下一个...")
                _interruptible_sleep(sleep_time, stop_event)

        logger.info(f"导出完成: 成功 {ok_count} 条, 失败 {fail_count} 条, "
                    f"跳过 {skipped_count} 条, AI过滤 {ai_skipped_count} 条")
        return ok_count, fail_count, skipped_count, ai_skipped_count

    def _find_weibo_dir(self, output_dir, weibo_id, export_format="md"):
        """查找指定微博ID导出文件所在的月份目录(在 年份/月份 目录树中)

        文件名格式: <博主>_<日期>_<微博ID>.md|docx
        返回目录路径;未找到返回 None
        """
        ext = ".docx" if export_format == "docx" else ".md"
        target = f"_{weibo_id}{ext}"
        try:
            for year_item in os.listdir(output_dir):
                year_path = os.path.join(output_dir, year_item)
                if not (os.path.isdir(year_path) and year_item.endswith("年")):
                    continue
                for month_item in os.listdir(year_path):
                    month_path = os.path.join(year_path, month_item)
                    if not (os.path.isdir(month_path) and month_item.endswith("月")):
                        continue
                    for f in os.listdir(month_path):
                        if f.endswith(target):
                            return month_path
        except Exception as e:
            logger.warning(f"查找微博文件时出错: {e}")
        return None

    def _weibo_file_exists(self, output_dir, weibo_id, export_format="md"):
        """检查指定微博ID是否已有导出文件"""
        return self._find_weibo_dir(output_dir, weibo_id, export_format) is not None

    @staticmethod
    def _mark_no_media(month_dir, weibo_id, media_type):
        """写"已确认无该类型媒体"标记文件,避免后续重复抓详情检查"""
        try:
            marker = os.path.join(month_dir, f"{weibo_id}.{media_type}.nomedia")
            with open(marker, "w", encoding="utf-8") as f:
                f.write("confirmed no media\n")
        except Exception as e:
            logger.warning(f"写无媒体标记失败: {e}")

    @staticmethod
    def _media_exists(month_dir, weibo_id, media_type="images"):
        """检查指定微博在该目录下是否已有媒体文件(如 images/xxx_1.jpg)

        存在标记文件 <weibo_id>.nomedia 时视为"已确认无该类型媒体",返回 True
        (避免对无图微博重复抓详情检查)
        """
        marker = os.path.join(month_dir, f"{weibo_id}.{media_type}.nomedia")
        if os.path.exists(marker):
            return True
        sub = os.path.join(month_dir, media_type)
        if not os.path.isdir(sub):
            return False
        prefix = f"{weibo_id}_"
        for f in os.listdir(sub):
            if f.startswith(prefix):
                return True
        return False

    @staticmethod
    def _publish_month_dir(publish_time, output_dir):
        """根据发布时间解析 年份年/月份 目录(如 2024年/10月)"""
        try:
            s = (publish_time or "").strip()
            # 支持 "26-4-7 21:58" / "2026-04-07" / "4月7日" 等格式
            m = re.match(r"(\d{2,4})-(\d{1,2})-\d{1,2}", s)
            if m:
                year = int(m.group(1))
                if year < 100:
                    year += 2000
                return os.path.join(output_dir, f"{year}年", f"{int(m.group(2))}月")
            m2 = re.search(r"(\d{4})年(\d{1,2})月", s)
            if m2:
                return os.path.join(output_dir, f"{int(m2.group(1))}年", f"{int(m2.group(2))}月")
        except Exception:
            pass
        return output_dir

    @staticmethod
    def _write_docx(detail, output_path, media_dir=None,
                    local_images=None, local_videos=None,
                    download_images=False, download_videos=False):
        """将详情导出为 docx 文件(使用 python-docx)

        - 全文使用宋体
        - 已下载的图片: 直接嵌入文档(按顺序)
        - 未下载的图片: 写入图片URL,并提示可勾选"下载图片"后重新导出
        - 视频: docx 无法嵌入,写入本地路径/链接并提示
        """
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn

        doc = Document()

        # 设置默认字体为宋体(含中文 eastAsia 字体)
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        def add_para(text, bold=False):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.bold = bold
            return p

        # URL 行: 带可点击超链接(蓝色下划线),查看原文章/评论区可直接点击
        p_url = doc.add_paragraph()
        run = p_url.add_run("URL：")
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 注意: 本方法是 @staticmethod,不能使用 self,用类名调用
        WeiboPCCrawler._add_hyperlink(p_url, detail['url'], detail['url'])
        add_para(f"发布时间：{detail['publish_time']}")
        add_para(f"词条：{detail['topics']}")
        add_para(f"正文字数：{len(detail['content'])}字符")
        add_para("—" * 20)
        add_para("正文：", bold=True)
        for line in detail['content'].splitlines():
            add_para(line)
        add_para("—" * 20)

        # 图片: 已下载则嵌入,否则写URL
        images = detail.get("images") or []
        if images:
            add_para("图片：", bold=True)
            local_images = local_images or []
            for idx, img_url in enumerate(images, 1):
                local_path = None
                if idx <= len(local_images) and local_images[idx - 1]:
                    local_path = local_images[idx - 1]
                if local_path and media_dir:
                    full = os.path.join(media_dir, local_path)
                    if os.path.exists(full):
                        try:
                            doc.add_picture(full, width=Cm(12))
                            continue
                        except Exception as e:
                            logger.warning(f"docx 插入图片失败 {full}: {e}")
                    add_para(f"图片{idx}: {local_path}(文件缺失)")
                else:
                    add_para(f"图片{idx}(未下载): {img_url}")
            if not download_images:
                add_para("提示: 如需将图片嵌入本文档,请在爬取时勾选“下载图片”后重新导出。")

        # 视频: docx 无法嵌入,写链接
        videos = detail.get("videos") or []
        if videos:
            add_para("视频：", bold=True)
            local_videos = local_videos or []
            for i2, v_url in enumerate(videos):
                if i2 < len(local_videos) and local_videos[i2]:
                    add_para(f"视频{i2 + 1}(已下载): {local_videos[i2]}")
                else:
                    add_para(f"视频{i2 + 1}: {v_url}")
            if not download_videos:
                add_para("提示: docx 无法嵌入视频,已下载的视频文件在同目录 videos/ 文件夹中;"
                         "如需下载请勾选“下载视频”后重新导出。")

        add_para("—" * 20)
        add_para(f"转发数：{detail['repost_count']}")
        add_para(f"评论数：{detail['comment_count']}")
        add_para(f"点赞数：{detail['like_count']}")
        add_para(f"保存时间：{detail['save_time']}")
        if detail.get("ai_quality") is not None:
            add_para(f"高质量可信度：{detail['ai_quality']}%")
        doc.save(output_path)

    @staticmethod
    def _copy_to_ai_dir(output_path, user_name, save_time, ai_title, ai_root,
                        copy_media=True, detail=None):
        """把刚导出的文章复制到 AI分类/AI_<博主>_高质量/ 并重命名为 <日期>_<AI标题>

        保留 DataPC 原文件(微博ID/媒体去重均依赖原文件)。
        copy_media=True(默认): 同时把该微博的 images/videos 一并复制到AI分类目录
        (修复: 此前只复制文章文件,导致AI副本里的图片引用是坏的)。
        copy_media=False: 不复制媒体,并把 md 副本中的本地媒体引用改写为
        原微博网络链接(不占空间,联网可看图);docx 图片已内嵌,直接复制即可。
        """
        try:
            dst_root = os.path.join(ai_root, f"AI_{user_name}_高质量")
            month_dir = WeiboPCCrawler._publish_month_dir(save_time, dst_root)
            os.makedirs(month_dir, exist_ok=True)
            date_part = ((save_time or "").split()[0]
                         if save_time and save_time != "未知时间" else "unknown")
            safe_title = re.sub(r'[\\/:*?"<>|\r\n]+', "", ai_title or "").strip() or "无标题"
            safe_title = safe_title[:40]
            ext = os.path.splitext(output_path)[1] or ".md"
            dst = os.path.join(month_dir, f"{date_part}_{safe_title}{ext}")
            if os.path.exists(dst):
                stem, e = os.path.splitext(dst)
                n = 2
                while os.path.exists(f"{stem}_{n}{e}"):
                    n += 1
                dst = f"{stem}_{n}{e}"

            if copy_media:
                shutil.copy2(output_path, dst)
                # 连带复制该微博的图片/视频(与事后AI分类行为一致)
                WeiboPCCrawler._copy_ai_media(output_path, month_dir)
            elif ext.lower() == ".md":
                # 不复制媒体: 把 md 中的本地媒体引用改写为原网络链接,避免破图
                try:
                    with open(output_path, "r", encoding="utf-8") as f:
                        md_text = f.read()
                    md_text = WeiboPCCrawler._rewrite_media_refs(md_text, detail)
                    with open(dst, "w", encoding="utf-8") as f:
                        f.write(md_text)
                except Exception as e:
                    logger.warning(f"改写AI副本媒体引用失败,直接复制原文件: {e}")
                    shutil.copy2(output_path, dst)
            else:
                # docx: 图片已内嵌在文档中,直接复制即可
                shutil.copy2(output_path, dst)

            logger.info(f"已复制到AI分类目录: {dst}")
            return dst
        except Exception as e:
            logger.warning(f"复制到AI分类目录失败: {e}")
            return None

    @staticmethod
    def _copy_ai_media(src_file, dst_dir):
        """把与文章同目录 images/ videos/ 下属于该微博的媒体复制到 AI 分类目录"""
        try:
            m = re.search(r"_([A-Za-z0-9]+)\.(?:md|docx)$", os.path.basename(src_file))
            if not m:
                return
            weibo_id = m.group(1)
            src_dir = os.path.dirname(src_file)
            for sub in ("images", "videos"):
                src_sub = os.path.join(src_dir, sub)
                if not os.path.isdir(src_sub):
                    continue
                dst_sub = os.path.join(dst_dir, sub)
                os.makedirs(dst_sub, exist_ok=True)
                prefix = f"{weibo_id}_"
                for fname in os.listdir(src_sub):
                    if fname.startswith(prefix):
                        s = os.path.join(src_sub, fname)
                        d = os.path.join(dst_sub, fname)
                        if os.path.abspath(s) != os.path.abspath(d):
                            shutil.copy2(s, d)
        except Exception as e:
            logger.warning(f"复制媒体到AI分类目录失败 {src_file}: {e}")

    @staticmethod
    def _rewrite_media_refs(md_text, detail):
        """把 md 中的本地媒体引用改写为原微博网络链接

        匹配行: ![图片N](本地路径) / [视频](本地路径)
        网络链接行(![图片N](https://...) / [视频链接](https://...))不受影响。
        """
        if not md_text or not detail:
            return md_text
        images = detail.get("images") or []
        videos = detail.get("videos") or []

        def _img_repl(m):
            idx = int(m.group(1))
            if 1 <= idx <= len(images) and images[idx - 1]:
                return f"![图片{idx}]({images[idx - 1]})"
            return m.group(0)

        viter = iter(videos)

        def _vid_repl(m):
            try:
                url = next(viter)
            except StopIteration:
                return m.group(0)
            return f"[视频]({url})" if url else m.group(0)

        try:
            md_text = re.sub(r"!\[图片(\d+)\]\([^)]+\)", _img_repl, md_text)
            md_text = re.sub(r"\[视频\]\([^)]+\)", _vid_repl, md_text)
        except Exception as e:
            logger.warning(f"改写媒体引用失败: {e}")
        return md_text

    @staticmethod
    def _add_hyperlink(paragraph, url, text):
        """在 docx 段落中添加可点击超链接(python-docx 原生不支持,手动写 XML)"""
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            part = paragraph.part
            r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            fonts = OxmlElement('w:rFonts')
            fonts.set(qn('w:eastAsia'), '宋体')
            rPr.append(fonts)
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rPr.append(rStyle)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0563C1')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            new_run.append(t)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
            return True
        except Exception as e:
            logger.warning(f"添加超链接失败: {e}")
            return False

    @staticmethod
    def _download_media(urls, base_dir, sub_dir, weibo_id):
        """下载图片/视频到 base_dir/sub_dir 下(带微博 Referer 防盗链)

        返回下载成功的本地相对路径列表(如 "images/xxx_1.jpg"),失败项为 None
        """
        try:
            import requests
        except ImportError:
            logger.warning("未安装 requests,无法下载媒体文件")
            return []
        target = os.path.join(base_dir, sub_dir)
        os.makedirs(target, exist_ok=True)
        headers = {
            "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                           "AppleWebKit/537.36 (KHTML, like Gecko) "
                           "Chrome/120.0.0.0 Safari/537.36 Edg/120.0.0.0"),
            "Referer": "https://weibo.com/",
        }
        results = []
        for idx, url in enumerate(urls, 1):
            try:
                ext = ".jpg"
                for cand in (".png", ".gif", ".webp", ".mp4"):
                    if cand in url.lower():
                        ext = cand
                        break
                fname = f"{weibo_id}_{idx}{ext}"
                fpath = os.path.join(target, fname)
                if not os.path.exists(fpath):
                    resp = requests.get(url, headers=headers, timeout=30)
                    if resp.status_code == 200:
                        with open(fpath, "wb") as f:
                            f.write(resp.content)
                        logger.info(f"已下载媒体: {os.path.join(sub_dir, fname)}")
                    else:
                        logger.warning(f"下载媒体失败 HTTP {resp.status_code}: {url}")
                        results.append(None)
                        continue
                results.append(os.path.join(sub_dir, fname).replace("\\", "/"))
            except Exception as e:
                logger.warning(f"下载媒体出错 {url}: {e}")
                results.append(None)
        return results

    def get_user_weibo_list(self, user_id, is_ori=1, is_forward=1, is_text=1, is_pic=1,
                            is_video=1, is_music=1, keyword="", start_date=None, end_date=None,
                            max_count=50, pause=False):
        """兼容旧版 weibo_selenium_PC.py 的接口:获取用户微博列表

        与 collect_weibo_ids 等价;pause=True 时在收集完成后暂停等待用户确认
        """
        if not self.check_login_status():
            logger.warning("未检测到登录状态,尝试手动登录")
            if not self.manual_login():
                logger.error("登录失败,无法继续")
                return None, []

        search_url = self.build_search_url(
            user_id, is_ori, is_forward, is_text, is_pic,
            is_video, is_music, keyword, start_date, end_date)
        logger.info(f"正在访问URL: {search_url}")
        self.driver.get(search_url)
        time.sleep(3)
        self.click_search_button()
        time.sleep(5)

        username = self.get_username(user_id, fallback_name=user_id)

        card_cls = self.resolve_card_class(user_id)
        time_cls = self.resolve_time_class(user_id)
        weibo_ids = self._scroll_and_collect(user_id, card_cls, time_cls, max_count)

        if pause:
            if self.wait_callback:
                self.wait_callback("程序已暂停,请检查浏览器中的页面状态,然后点击「确定」继续...")
            else:
                input("程序已暂停,请检查浏览器中的页面状态。按回车键继续或关闭浏览器...")

        return username, weibo_ids

    # ---------- 文件辅助 ----------

    def save_weibo_ids_to_file(self, weibo_ids, filename):
        """将微博 ID 保存到文本文件,用逗号分隔"""
        try:
            os.makedirs(os.path.dirname(filename), exist_ok=True)
            with open(filename, "w", encoding="utf-8") as f:
                f.write(",".join(weibo_ids))
            logger.info(f"微博ID已保存到文件: {filename}")
            return True
        except Exception as e:
            logger.error(f"保存微博ID到文件时出错: {e}")
            return False

    def close(self, force_close=True):
        """关闭浏览器"""
        try:
            if force_close:
                self.driver.quit()
                logger.info("浏览器已关闭")
        except Exception:
            pass


# ---------------------------------------------------------------------------
# 微博状态记录: 每个 ID 的导出状态/AI判定/转评赞/媒体
# ---------------------------------------------------------------------------

STATS_FILE = "微博状态.json"


class WeiboStatsStore:
    """博主级状态记录(DataPC/<博主>_<ID>/微博状态.json)

    记录每个微博 ID 的:
      exported       是否已导出文件
      file           导出文件名(所在月份目录)
      publish        发布时间(用于定位月份目录)
      ai             {quality, checked_at} AI判定分数与时间
      stats          {repost, comment, like, checked_at} 转评赞与最后检查时间
      media          {images, videos} 是否已下载

    用途:
      1. 重跑时复用 AI 判定结果,避免重复抓详情/调用大模型(修复 AI 评分漂移问题)
      2. 支撑"更新转评赞":按已导出 ID 逐篇抓详情,只刷新数字不重写正文
      3. 后续(如 V0.7.0 增量模式)可直接读取状态判断"哪些文章还没处理"
    """

    def __init__(self, blogger_dir):
        self.path = os.path.join(blogger_dir, STATS_FILE)
        self.data = {"records": {}}
        self.load()

    def load(self):
        try:
            if os.path.exists(self.path):
                with open(self.path, "r", encoding="utf-8") as f:
                    d = json.load(f)
                if isinstance(d, dict) and isinstance(d.get("records"), dict):
                    self.data = d
        except Exception as e:
            logger.warning(f"读取状态记录失败: {e}")

    def save(self):
        try:
            with open(self.path, "w", encoding="utf-8") as f:
                json.dump(self.data, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            logger.warning(f"写入状态记录失败: {e}")
            return False

    def get(self, weibo_id):
        return self.data["records"].get(weibo_id)

    def set_record(self, weibo_id, **fields):
        rec = self.data["records"].setdefault(weibo_id, {})
        for k, v in fields.items():
            if v is None:
                rec.pop(k, None)
            else:
                rec[k] = v

    def exported_ids(self):
        return [wid for wid, r in self.data["records"].items()
                if r.get("exported") and r.get("file")]

    def exported_ids_in_range(self, start_date=None, end_date=None):
        """按 publish 过滤已导出ID(YYYY-MM-DD;publish 未知的保守包含)"""
        start = datetime.strptime(start_date, "%Y-%m-%d") if start_date else None
        end = datetime.strptime(end_date, "%Y-%m-%d") if end_date else None
        if start is None and end is None:
            return self.exported_ids()
        out = []
        for wid in self.exported_ids():
            d = _parse_publish_date((self.get(wid) or {}).get("publish"))
            if d is None:
                out.append(wid)  # 未知日期:保守包含,确保不漏
                continue
            if start and d < start:
                continue
            if end and d > end:
                continue
            out.append(wid)
        return out


def _now_str():
    return datetime.now().strftime("%y-%m-%d %H:%M")


def _update_stats_in_file(path, repost, comment, like):
    """仅更新 md/docx 文件中的 转发数/评论数/点赞数 三行(不重写正文)

    返回 True=更新成功, False=读取/写入失败
    """
    try:
        if path.endswith(".md"):
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
            text = re.sub(r"(>转发数[:：]\s*)[\d.,万亿]+", rf"\g<1>{repost}", text)
            text = re.sub(r"(>评论数[:：]\s*)[\d.,万亿]+", rf"\g<1>{comment}", text)
            text = re.sub(r"(>点赞数[:：]\s*)[\d.,万亿]+", rf"\g<1>{like}", text)
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)
            return True
        if path.endswith(".docx"):
            from docx import Document
            doc = Document(path)
            updated = False
            for p in doc.paragraphs:
                m = re.match(r"^(转发数|评论数|点赞数)[:：]\s*[\d.,万亿]+$", p.text.strip())
                if not m:
                    continue
                label, value = m.group(1), {"转发数": repost, "评论数": comment, "点赞数": like}[m.group(1)]
                if not p.runs:
                    continue
                p.runs[0].text = f"{label}：{value}"
                for run in p.runs[1:]:
                    run.text = ""
                updated = True
            if updated:
                doc.save(path)
                return True
            logger.warning(f"docx 中未找到转评赞行: {path}")
            return False
    except Exception as e:
        logger.warning(f"更新转评赞失败 {path}: {e}")
        return False


def _parse_publish_date(publish):
    """从状态记录 publish 字段解析日期,兼容 "25-1-10 12:00" / "2025-01-10" 等格式"""
    if not publish:
        return None
    m = re.search(r"(\d{2,4})-(\d{1,2})-(\d{1,2})", str(publish))
    if not m:
        return None
    y = int(m.group(1))
    y = y + 2000 if y < 100 else y
    try:
        return datetime(y, int(m.group(2)), int(m.group(3)))
    except ValueError:
        return None


def update_stats_for_blogger(user_id, user_name, blogger_dir, headless=False,
                             user_data_dir=None, min_interval=3, max_interval=8,
                             progress_callback=None, wait_callback=None,
                             stop_event=None, keep_browser_open=False,
                             start_date=None, end_date=None):
    """按状态记录更新已导出文章的转评赞(月末/季末/年末跑一次)

    start_date/end_date(YYYY-MM-DD)可选: 只更新该日期范围内的已导出文章
    (按状态记录 publish 过滤;publish 未知的记录在给定范围时保守包含,确保不漏)。

    逐篇抓详情 → 只更新 md/docx 中的 转发/评论/点赞 数字与状态记录检查时间;
    不重写正文、不调用 AI。支持优雅停止(每篇完成即写状态,中断不丢已更新部分)。
    返回 {"updated": n, "failed": n, "message": str}
    """
    ensure_logger()
    store = WeiboStatsStore(blogger_dir)
    ids = store.exported_ids_in_range(start_date, end_date)
    if not ids:
        msg = ("状态记录中没有符合范围要求的已导出文章"
               + (f"({start_date or '全部'} ~ {end_date or '全部'})"
                  if (start_date or end_date) else "(请先爬取一次生成状态记录)"))
        logger.warning(msg)
        return {"updated": 0, "failed": 0, "message": msg}

    crawler = WeiboPCCrawler(headless=headless, user_data_dir=user_data_dir,
                             wait_callback=wait_callback)
    updated = failed = 0
    try:
        if not crawler.check_login_status():
            logger.warning("未检测到登录状态,尝试手动登录")
            if not crawler.manual_login():
                return {"updated": 0, "failed": 0, "message": "登录失败,无法更新转评赞"}

        total = len(ids)
        for i, wid in enumerate(ids, 1):
            if stop_event is not None and stop_event.is_set():
                logger.info("收到停止请求,停止更新转评赞")
                break
            if progress_callback:
                try:
                    progress_callback(i, total, wid)
                except Exception:
                    pass
            detail = crawler.get_weibo_detail(wid, user_id)
            if not detail:
                failed += 1
                logger.warning(f"获取详情失败,跳过: {wid}")
                continue
            rec = store.get(wid) or {}
            fname = rec.get("file") or ""
            path = ""
            if fname:
                # 优先按记录中的月份目录定位,找不到再全目录搜索
                month_dir = WeiboPCCrawler._publish_month_dir(
                    rec.get("publish") or "", blogger_dir)
                cand = os.path.join(month_dir, fname)
                if os.path.exists(cand):
                    path = cand
            if not path:
                found = crawler._find_weibo_dir(blogger_dir, wid)
                if found:
                    for f in os.listdir(found):
                        if f.endswith(f"_{wid}.md") or f.endswith(f"_{wid}.docx"):
                            path = os.path.join(found, f)
                            break
            if path and _update_stats_in_file(path, detail["repost_count"],
                                              detail["comment_count"],
                                              detail["like_count"]):
                updated += 1
                logger.info(f"已更新转评赞: {os.path.basename(path)}")
            else:
                failed += 1
                logger.warning(f"未找到文章文件,跳过: {wid}")
            store.set_record(
                wid,
                stats={"repost": detail["repost_count"], "comment": detail["comment_count"],
                       "like": detail["like_count"], "checked_at": _now_str()})
            store.save()
            if i < total:
                _interruptible_sleep(random.randint(min_interval, max_interval), stop_event)
    finally:
        if not keep_browser_open:
            crawler.close(force_close=True)

    return {"updated": updated, "failed": failed, "message": ""}


# ---------------------------------------------------------------------------
# V0.6.0 导出到笔记库(Obsidian / 通用 Markdown 文件夹)
# 设计:CONTEXT.md + docs/adr/0001-0002 + 产品设计文档 §10
# ---------------------------------------------------------------------------

NOTE_ROOT = "微博长文"          # 目标库内笔记根目录
NOTE_ATTACH_DIR = "附件"        # 位于 NOTE_ROOT 下: 微博长文/附件/<博主>/<年>/<月>/
NOTE_CSS = """/* 微博长文收集记录整理器 - 笔记样式 */
/* 启用后:阅读/编辑视图默认隐藏笔记属性面板(需要时 Ctrl+P → Toggle properties 查看) */
.metadata-container,
.markdown-reading-view .metadata-container,
.markdown-preview-view .metadata-container,
.markdown-source-view .metadata-container {
    display: none !important;
}
"""

# 属性键顺序(英文键 + 中文值;空值字段省略)
NOTE_ATTRIBUTE_KEYS = (
    "title", "author", "source_url", "weibo_id", "publish_date",
    "topics", "category", "quality", "repost", "comment", "like",
    "imported_at", "tags",
)


def safe_note_title(text, max_len=40):
    """清洗用于笔记文件名/标题的文本(Windows 非法字符 + 限长)"""
    return re.sub(r'[\\/:*?"<>|\r\n]+', "", text or "").strip()[:max_len] or "无标题"


def publish_to_iso(publish):
    """任意 publish 格式 -> 'YYYY-MM-DD';无法解析返回 ''"""
    d = _parse_publish_date(publish)
    return d.strftime("%Y-%m-%d") if d else ""


def _extract_source_md(path):
    """从源文章 md 解析笔记所需字段;docx 提取正文文本

    返回 dict:url/publish/topics/body/images/videos/repost/comment/like/quality
    """
    result = {"url": "", "publish": "", "topics": "", "body": "",
              "images": [], "videos": [], "repost": None, "comment": None,
              "like": None, "quality": None, "saved_at": ""}
    try:
        if path.endswith(".md"):
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
            m = re.search(r">URL[：:]\s*\[[^\]]*\]\(([^)]+)\)|>URL[：:]\s*(\S+)", text)
            if m:
                result["url"] = m.group(1) or m.group(2) or ""
            m = re.search(r">发布时间[：:]\s*(\S+\s*\S*)", text)
            if m:
                result["publish"] = m.group(1).strip()
            m = re.search(r">词条[：:]\s*(.+)", text)
            if m:
                result["topics"] = m.group(1).strip()
            m = re.search(r"正文[：:]\s*\n(.*?)\n\s*---", text, re.DOTALL)
            if m:
                result["body"] = m.group(1).strip()
            for label, key in (("转发数", "repost"), ("评论数", "comment"), ("点赞数", "like")):
                m = re.search(rf">\s*{label}[：:]\s*([\d.,万亿]+)", text)
                if m:
                    result[key] = m.group(1)
            m = re.search(r">高质量可信度[：:]\s*([\d.]+)%?", text)
            if m:
                result["quality"] = m.group(1)
            m = re.search(r">保存时间[：:]\s*(\S+\s*\S*)", text)
            if m:
                result["saved_at"] = m.group(1).strip()
            # 图片行:![图片N](path) → (显示序号, 引用)
            for m in re.finditer(r"!\[图片\d+\]\(([^)]+)\)", text):
                result["images"].append(m.group(1).strip())
            # 视频行:[视频](path) / [视频链接](url)
            for m in re.finditer(r"\[视频[链接]*\]\(([^)]+)\)", text):
                result["videos"].append(m.group(1).strip())
        elif path.endswith(".docx"):
            from docx import Document
            doc = Document(path)
            lines = [p.text.strip() for p in doc.paragraphs]
            in_body = False
            parts = []
            for t in lines:
                m = re.match(r"^(URL|发布时间|词条)[：:]\s*(.+)$", t)
                if m:
                    if m.group(1) == "URL":
                        result["url"] = m.group(2)
                    elif m.group(1) == "发布时间":
                        result["publish"] = m.group(2)
                    else:
                        result["topics"] = m.group(2)
                m2 = re.match(r"^(转发数|评论数|点赞数)[：:]\s*([\d.,万亿]+)$", t)
                if m2:
                    key = {"转发数": "repost", "评论数": "comment", "点赞数": "like"}[m2.group(1)]
                    result[key] = m2.group(2)
                if t.strip() == "正文：":
                    in_body = True
                    continue
                if in_body and t.startswith("—"):
                    break
                if in_body:
                    parts.append(t)
            result["body"] = "\n".join(parts).strip()
    except Exception as e:
        logger.warning(f"解析源文章失败 {path}: {e}")
    return result


def iter_source_articles(blogger_dir, source="data", start_date=None, end_date=None,
                         stats_store=None, stop_event=None, author=None,
                         ai_root=None):
    """枚举博主源文章(AI分类 或 DataPC)

    source="ai"   : AI分类/AI_<博主>_<类别>/<年>/<月>/*.md|docx,类别取自目录名
    source="data" : DataPC/<博主>_<ID>/<年>/<月>/<博主>_<日期>_<ID>.md|docx
    每项 dict: author/weibo_id/publish/publish_date/title/category/quality/
               ai_title/topics/stats/body/images/videos/source_path
    """
    basename = os.path.basename(blogger_dir)
    # author 缺省时从目录名推导(去掉尾部 _ID)
    if not author:
        m = re.match(r"^(.+)_(\d+)$", basename)
        author = m.group(1) if m else basename
    start = datetime.strptime(start_date, "%Y-%m-%d") if start_date else None
    end = datetime.strptime(end_date, "%Y-%m-%d") if end_date else None

    def _in_range(pub_iso):
        if not pub_iso:
            return True
        d = datetime.strptime(pub_iso, "%Y-%m-%d")
        if start and d < start:
            return False
        if end and d > end:
            return False
        return True

    def _clean_topics(topics_text):
        """词条首项并清洗 # 号(用作标题兜底)"""
        for t in re.split(r"[、，,;；\s]+", topics_text or ""):
            t = t.strip().strip("#").strip()
            if t and len(t) <= 40:
                return t
        return ""

    def _yield(fpath, category, ai_title, fdate):
        if stop_event is not None and stop_event.is_set():
            return
        meta = _extract_source_md(fpath)
        pub_iso = publish_to_iso(meta.get("publish") or "") or (
            fdate.strftime("%Y-%m-%d") if fdate else "")
        if not _in_range(pub_iso):
            return
        wid_m = re.search(r"_([A-Za-z0-9]+)\.(?:md|docx)$", os.path.basename(fpath))
        weibo_id = wid_m.group(1) if wid_m else ""
        # 标题优先级: AI标题(分类文件名) > 话题首项 > 无标题
        title = ai_title or _clean_topics(meta.get("topics") or "") or "无标题"
        yield {
            "author": author, "weibo_id": weibo_id,
            "publish": meta.get("publish") or "", "publish_date": pub_iso,
            "title": title, "category": category,
            "quality": meta.get("quality"), "ai_title": ai_title or "",
            "topics": meta.get("topics") or "", "body": meta.get("body") or "",
            "images": meta.get("images") or [], "videos": meta.get("videos") or [],
            "stats": {"repost": meta.get("repost"), "comment": meta.get("comment"),
                      "like": meta.get("like")},
            "url": meta.get("url") or "", "saved_at": meta.get("saved_at") or "",
            "source_path": fpath,
        }

    if source in ("ai", "ai_high"):
        # AI 分类目录位于 AI分类/AI_<博主>_<类别>/(与 DataPC 同级,不在 blogger_dir 内)
        # source="ai_high" = 仅"高质量"类别;source="ai" = 全部类别
        base = ai_root or os.path.join(app_dir(), "AI分类")
        if not os.path.isdir(base):
            logger.warning(f"AI分类目录不存在: {base}")
            return
        for cat_name in sorted(os.listdir(base)):
            if not cat_name.startswith(f"AI_{author}_"):
                continue
            category = "AI未分类"
            for suffix in ("高质量", "广告", "可疑", "其他低质量"):
                if cat_name.endswith(f"_{suffix}"):
                    category = suffix
                    break
            if category == "AI未分类":
                continue
            if source == "ai_high" and category != "高质量":
                continue
            cat_dir = os.path.join(base, cat_name)
            for year_name in sorted(os.listdir(cat_dir)):
                ypath = os.path.join(cat_dir, year_name)
                if not (year_name.endswith("年") and os.path.isdir(ypath)):
                    continue
                for month_name in sorted(os.listdir(ypath)):
                    mpath = os.path.join(ypath, month_name)
                    if not (month_name.endswith("月") and os.path.isdir(mpath)):
                        continue
                    for f in sorted(os.listdir(mpath)):
                        if not f.endswith((".md", ".docx")):
                            continue
                        # AI 文件名: 日期_标题 → 标题段
                        dm = re.match(r"^\d{2,4}-\d{1,2}-\d{1,2}_(.+)\.(?:md|docx)$", f)
                        ai_title = dm.group(1) if dm else ""
                        yield from _yield(os.path.join(mpath, f), category,
                                          ai_title, None)
    else:
        for year_name in sorted(os.listdir(blogger_dir)):
            ypath = os.path.join(blogger_dir, year_name)
            if not (year_name.endswith("年") and os.path.isdir(ypath)):
                continue
            for month_name in sorted(os.listdir(ypath)):
                mpath = os.path.join(ypath, month_name)
                if not (month_name.endswith("月") and os.path.isdir(mpath)):
                    continue
                for f in sorted(os.listdir(mpath)):
                    if not f.endswith((".md", ".docx")):
                        continue
                    # DataPC 文件名: <博主>_<日期>_<ID>;日期用于兜底
                    fm = re.search(r"_(\d{1,4})-(\d{1,2})-(\d{1,2})_", f)
                    fdate = None
                    if fm:
                        y = int(fm.group(1))
                        if y < 100:
                            y += 2000
                        try:
                            fdate = datetime(y, int(fm.group(2)), int(fm.group(3)))
                        except ValueError:
                            fdate = None
                    ai_title = None
                    if stats_store:
                        wid_m = re.search(r"_([A-Za-z0-9]+)\.(?:md|docx)$", f)
                        if wid_m:
                            rec = stats_store.get(wid_m.group(1)) or {}
                            ai_title = rec.get("ai_title") or None
                    yield from _yield(os.path.join(mpath, f), "AI未分类",
                                      ai_title, fdate)


def build_note_markdown(article, imported_at=None, show_meta=True):
    """渲染笔记 md:13 字段 frontmatter + 正文 + 图片/视频引用

    show_meta=True 时正文前后带关键信息引用块(URL/时间/词条/字数 与 转评赞),
    与手工导入的旧格式一致;属性(frontmatter)始终保留供机器读取
    """
    imported_at = imported_at or _now_str()
    tags = [article["author"]]
    if article.get("category") and article["category"] != "AI未分类":
        tags.append(article["category"])
    if article.get("publish_date") and len(article["publish_date"]) >= 7:
        tags.append(article["publish_date"][:7])
    # 话题标签:去 #、按分隔符与空格拆分、限长、去重
    for topic in re.split(r"[、，,;；\s]+", article.get("topics") or ""):
        topic = topic.strip().strip("#").strip()
        if topic and len(topic) <= 20 and topic not in tags:
            tags.append(topic)

    attrs = {
        "title": article.get("title") or "无标题",
        "author": article.get("author") or "",
        "source_url": article.get("source_url") or "",
        "weibo_id": article.get("weibo_id") or "",
        "publish_date": article.get("publish_date") or "",
        "topics": article.get("topics") or "",
        "category": article.get("category") or "",
        "quality": article.get("quality"),
        "repost": (article.get("stats") or {}).get("repost"),
        "comment": (article.get("stats") or {}).get("comment"),
        "like": (article.get("stats") or {}).get("like"),
        "imported_at": imported_at,
        "tags": tags,
    }
    lines = ["---"]
    for key in NOTE_ATTRIBUTE_KEYS:
        val = attrs.get(key)
        if val in (None, "", [], {}):
            continue
        if key == "tags":
            rendered = ", ".join(f'"{t}"' for t in val)
            lines.append(f"tags: [{rendered}]")
        else:
            lines.append(f"{key}: {val}")
    lines.append("---")
    lines.append("")
    # 阅读从文件标题 + 关键信息块直接开始;不再加"## 正文"标题(与笔记标题重复)
    if show_meta:
        line_meta = []
        if article.get("url"):
            line_meta.append(f"URL：{article['url']}")
        if article.get("publish"):
            line_meta.append(f"发布时间：{article['publish']}")
        if article.get("topics"):
            line_meta.append(f"词条：{article['topics']}")
        if article.get("body"):
            line_meta.append(f"正文字数：{len(article['body'])}字符")
        for lm in line_meta:
            lines.append(f"> {lm}")
        if line_meta:
            lines.append("")
    lines.append(article.get("body") or "")
    if show_meta:
        stats = article.get("stats") or {}
        line_stats = []
        if stats.get("like") is not None:
            line_stats.append(f"点赞数：{stats['like']}")
        if stats.get("comment") is not None:
            line_stats.append(f"评论数：{stats['comment']}")
        if stats.get("repost") is not None:
            line_stats.append(f"转发数：{stats['repost']}")
        if article.get("saved_at"):
            line_stats.append(f"保存时间：{article['saved_at']}")
        if article.get("quality") is not None:
            line_stats.append(f"高质量可信度：{article['quality']}%")
        if line_stats:
            lines.append("")
            lines.append("---")
            lines.append("")
            for ls in line_stats:
                lines.append(f"> {ls}")
    if article.get("images"):
        lines.append("")
        lines.append("### 图片")
        for ref in article["images"]:
            lines.append(f"![图片]({ref})")
    if article.get("videos"):
        lines.append("")
        lines.append("### 视频")
        for ref in article["videos"]:
            lines.append(f"[视频]({ref})")
    return "\n".join(lines)


def scan_target_notes(target_root, author):
    """扫描目标库 微博长文/<author> 下所有笔记,解析 (weibo_id, publish_date) 幂等键

    返回 {(weibo_id, publish_date): note_path}(支持笔记被改名/移动)
    """
    base = os.path.join(target_root, NOTE_ROOT, author)
    found = {}
    if not os.path.isdir(base):
        return found
    for year_name in sorted(os.listdir(base)):
        ypath = os.path.join(base, year_name)
        if not (year_name.endswith("年") and os.path.isdir(ypath)):
            continue
        for month_name in sorted(os.listdir(ypath)):
            mpath = os.path.join(ypath, month_name)
            if not (month_name.endswith("月") and os.path.isdir(mpath)):
                continue
            for f in sorted(os.listdir(mpath)):
                if not f.endswith(".md"):
                    continue
                fp = os.path.join(mpath, f)
                wid = pub = ""
                try:
                    with open(fp, "r", encoding="utf-8") as fh:
                        head = fh.read(2000)
                except Exception:
                    continue
                m = re.search(r"^weibo_id:\s*(\S+)$", head, re.M)
                if m:
                    wid = m.group(1)
                m = re.search(r"^publish_date:\s*(\S+)$", head, re.M)
                if m:
                    pub = m.group(1)
                if wid and pub:
                    found[(wid, pub)] = fp
    return found


def _resolve_note_image(ref, note_dir, attachment_dir, copy_images):
    """图片引用 -> 目标笔记内可用引用(URL 原样;本地文件复制到附件并改相对路径)"""
    if re.match(r"^https?://", ref):
        return ref, None
    src = os.path.abspath(ref)
    if not os.path.exists(src):
        return ref, None  # 本地文件缺失:保持原引用
    if not copy_images:
        return ref, None
    dst = os.path.join(attachment_dir, os.path.basename(src))
    try:
        if os.path.abspath(src) != os.path.abspath(dst):
            os.makedirs(attachment_dir, exist_ok=True)
            shutil.copy2(src, dst)
        return os.path.relpath(dst, note_dir).replace("\\", "/"), None
    except Exception as e:
        logger.warning(f"复制附件失败 {src}: {e}")
        return ref, None


def export_notes_for_blogger(user_name, user_id, blogger_dir, target_root,
                             source="ai", target="obsidian",
                             start_date=None, end_date=None,
                             copy_images=True, conflict="skip",
                             overwrite=False, progress_callback=None,
                             stop_event=None, ai_root=None, write_css=True,
                             show_meta=True):
    """把博主源文章导入目标笔记库(V0.6.0 导出主流程)

    参数:
      blogger_dir   DataPC/<博主>_<ID>(同时用于读取状态与定位 AI分类 同级约定)
      target_root   vault 或通用笔记库根目录
      source        "ai"=仅AI分类结果 | "data"=全部已导出
      target        "obsidian" | "markdown-folder"(栈中仅作标记;open 动作由上层处理)
      conflict      "skip" | "rename"(同名不同源:跳过或加后缀)
      overwrite     同源已存在时是否覆盖(默认 False 跳过)
    返回 {"imported": [...], "updated": n, "skipped": n, "conflicts": [...],
          "failed": [...], "planned": n}
    """
    ensure_logger()
    target_root = os.path.abspath(target_root)
    author = (user_name or "").strip() or os.path.basename(blogger_dir)
    stats_store = WeiboStatsStore(blogger_dir) if os.path.isdir(blogger_dir) else None

    # 幂等扫描:该博主已存在的笔记(可能是被改名/移动过的)
    existing = scan_target_notes(target_root, author)

    imported, conflicts, failed = [], [], []
    updated_n = skipped_n = 0
    total_planned = 0

    for article in iter_source_articles(blogger_dir, source=source,
                                        start_date=start_date, end_date=end_date,
                                        stats_store=stats_store,
                                        stop_event=stop_event,
                                        author=author, ai_root=ai_root):
        if stop_event is not None and stop_event.is_set():
            logger.info("收到停止请求,导出停止")
            break
        total_planned += 1
        if progress_callback:
            progress_callback(total_planned, 0, article.get("weibo_id") or "")
        key = (article.get("weibo_id") or "", article.get("publish_date") or "")
        if key in existing and not overwrite:
            skipped_n += 1
            continue

        # 目标路径: 微博长文/<author>/<年>/<月>/yy-m-d_标题.md
        pub = article.get("publish_date") or ""
        if len(pub) < 10:
            pub = datetime.now().strftime("%Y-%m-%d")
        note_dir = os.path.join(target_root, NOTE_ROOT, author,
                                f"{pub[:4]}年", f"{int(pub[5:7])}月")
        # 附件: 微博长文/附件/<author>/<年>/<月>/(与笔记同库同体系)
        note_attach = os.path.join(target_root, NOTE_ROOT, NOTE_ATTACH_DIR,
                                   author, f"{pub[:4]}年", f"{int(pub[5:7])}月")
        # 文件名与用户手工导入风格一致: 26-2-1_标题.md(两位年,月日不补零)
        short_date = f"{int(pub[2:4])}-{int(pub[5:7])}-{int(pub[8:10])}"
        stem = f"{short_date}_{safe_note_title(article.get('title'))}"
        filename = f"{stem}.md"
        note_path = os.path.join(note_dir, filename)
        if os.path.exists(note_path) and key not in existing:
            # 同名但不同源:按策略处理
            if conflict == "rename":
                n = 2
                while os.path.exists(os.path.join(note_dir, f"{stem}_{n}.md")):
                    n += 1
                note_path = os.path.join(note_dir, f"{stem}_{n}.md")
            else:
                conflicts.append(os.path.basename(article.get("source_path") or filename))
                continue

        try:
            os.makedirs(note_dir, exist_ok=True)
            # 源 md 中图片为相对源文章目录的路径 → 转绝对再解析
            abs_images = []
            for ref in article.get("images") or []:
                if re.match(r"^https?://", ref):
                    abs_images.append(ref)
                else:
                    abs_images.append(os.path.join(
                        os.path.dirname(article["source_path"]), ref))
            article["images"] = abs_images
            md_text = build_note_markdown(article, show_meta=show_meta)
            for ref in list(article.get("images") or []):
                new_ref, _ = _resolve_note_image(ref, note_dir, note_attach,
                                                 copy_images)
                md_text = md_text.replace(f"![图片]({ref})", f"![图片]({new_ref})", 1)
            with open(note_path, "w", encoding="utf-8") as f:
                f.write(md_text)
            if key in existing:
                updated_n += 1
            else:
                imported.append(os.path.relpath(note_path, target_root))
            # 状态记录标记"上次导入时间"(仅展示;幂等依据是目标笔记 frontmatter)
            if stats_store and article.get("weibo_id"):
                stats_store.set_record(article["weibo_id"],
                                       obsidian_imported_at=_now_str())
        except Exception as e:
            logger.warning(f"导出笔记失败 {article.get('source_path')}: {e}")
            failed.append(os.path.basename(article.get("source_path") or ""))

    return {"imported": imported, "updated": updated_n, "skipped": skipped_n,
            "conflicts": conflicts, "failed": failed,
            "planned": total_planned, "css_hint": _write_css_snippet(
                target_root, target, write_css)}


def _write_css_snippet(target_root, target, write_css):
    """向 vault 写入隐藏属性的 CSS 片段(仅 obsidian 目标;失败不影响导入)

    返回提示文本(空=未写入/失败)
    """
    if not write_css or target != "obsidian":
        return ""
    try:
        snippets = os.path.join(target_root, ".obsidian", "snippets")
        os.makedirs(snippets, exist_ok=True)
        with open(os.path.join(snippets, "weibo-notes.css"), "w",
                  encoding="utf-8") as f:
            f.write(NOTE_CSS)
        return ("\n提示:已写入 weibo-notes.css;如需阅读视图默认隐藏属性面板:"
                "Obsidian 设置→外观→CSS 片段→启用 weibo-notes"
                "(Ctrl+P → Toggle properties 可随时查看)")
    except Exception as e:
        logger.warning(f"写入CSS片段失败: {e}")
        return ""


# ---------------------------------------------------------------------------
# 一键任务:收集 + 导出
# ---------------------------------------------------------------------------

def run_task(user_id, user_name, start_date, end_date,
             headless=False, user_data_dir=None, max_count=500,
             keyword="", is_ori=1, is_forward=0, is_text=1, is_pic=1,
             is_video=1, is_music=1, manual_callback=None,
             wait_callback=None, progress_callback=None, data_root=None,
             keep_browser_open=False, skip_export=False,
             min_interval=3, max_interval=8,
             download_images=False, download_videos=False,
             export_format="md", skip_existing=False, min_words=0,
             ai_enabled=False, ai_config=None, ai_usage_callback=None,
             ai_rename=False, ai_root=None, ai_copy_media=True,
             detected_callback=None, stop_event=None, force_rejudge=False):
    """一键爬取任务:收集指定时间范围的微博ID并导出

    参数:
      user_id / user_name    博主微博ID与昵称
      start_date / end_date  时间范围(YYYY-MM-DD,含两端)
      headless               无头模式
      user_data_dir          Edge 用户数据目录(复用登录态)
      manual_callback        手动输入类名的回调,签名 (key, current) -> str
      wait_callback          等待用户操作的回调(如手动登录后继续),签名 (message) -> None
      progress_callback      进度回调(用于 GUI 显示)
      data_root              数据输出根目录,默认 DataPC
      keep_browser_open      完成后是否保留浏览器窗口
      skip_export            只收集ID,不导出
      min_interval/max_interval  每条微博之间随机等待秒数范围
      download_images/videos 是否下载图片/视频
      export_format          导出格式 "md" 或 "docx"
      skip_existing          跳过已存在同ID文件的微博(避免重复抓取)
      min_words              正文字数低于该值的文章不导出(0=不限制);
                             同时用于列表页粗过滤(低于该字数的微博不进详情页)
      ai_enabled             启用AI实时判断(需 ai_config 已配置API)
      ai_config              AIConfig 实例;为 None 时尝试读取 ai_config.json
      ai_usage_callback      token 用量回调(每次AI调用后调用)
      ai_rename              AI通过的文章同时生成标题并复制到 AI分类 目录
      ai_root                AI分类输出根目录(默认 程序目录/AI分类)
      ai_copy_media          AI分类副本是否连带复制图片/视频
                             (False 时md副本改为引用原网络链接)
      force_rejudge          重跑时对状态记录中已判定过的文章重新调用AI(默认复用上次分数)
      detected_callback      类名更新通知回调(同步GUI输入框用),签名 (key, value)
      stop_event             threading.Event,设置后优雅停止(当前条目收尾后停止)

    返回 dict:
      {"username", "weibo_ids", "txt_file", "md_dir", "exported", "failed",
       "skipped", "ai_skipped", "ai_tokens"}
    """
    ensure_logger()
    if not data_root:
        data_root = os.path.join(app_dir(), "DataPC")

    # 间隔过短提示(路线图要求)
    if min_interval < 2:
        logger.warning("提示:爬取间隔过短(<2秒)可能触发微博风控,建议设置 3 秒以上。")

    crawler = WeiboPCCrawler(headless=headless, user_data_dir=user_data_dir,
                             wait_callback=wait_callback,
                             manual_callback=manual_callback,
                             detected_callback=detected_callback)

    # 日期 +1 天:抵消微博时间换算导致的实际搜索范围偏移(沿用原脚本经验)
    start_search = (datetime.strptime(start_date, "%Y-%m-%d") + timedelta(days=1)).strftime("%Y-%m-%d")
    end_search = (datetime.strptime(end_date, "%Y-%m-%d") + timedelta(days=1)).strftime("%Y-%m-%d")

    result = {
        "username": user_name, "weibo_ids": [], "txt_file": None,
        "md_dir": None, "exported": 0, "failed": 0, "skipped": 0,
        "ai_skipped": 0, "ai_tokens": 0, "stopped": False,
    }

    # AI 实时判断: 构建分类器(API 未配置时仅警告,不启用AI)
    ai_classifier = None
    user_usage_callback = ai_usage_callback  # 外部传入的token回调
    ai_tokens = [0]  # 用列表累积,供回调读取
    if ai_enabled:
        try:
            from weibo_ai import AIConfig, AIClient, AIClassifier
            if ai_config is None:
                ai_config = AIConfig()
            if not ai_config.is_configured():
                logger.warning("启用了AI实时判断,但未配置API Key(ai_config.json),本次不启用AI")
            else:
                client = AIClient(ai_config.get("api_key", ""),
                                  ai_config.get("base_url", ""),
                                  ai_config.get("model", ""))
                ai_classifier = AIClassifier(client, ai_config)

                def _ai_usage_cb(delta):
                    ai_tokens[0] += int(delta or 0)
                    if user_usage_callback:
                        user_usage_callback(ai_tokens[0])
                ai_usage_callback = _ai_usage_cb
        except Exception as e:
            logger.warning(f"初始化AI分类器失败,本次不启用AI: {e}")
            ai_enabled = False

    try:
        # 1. 收集 ID(列表页按 min_words 粗过滤,低于字数的微博不进详情页)
        username, weibo_ids = crawler.collect_weibo_ids(
            user_id=user_id,
            start_date=start_search,
            end_date=end_search,
            max_count=max_count,
            keyword=keyword,
            is_ori=is_ori, is_forward=is_forward,
            is_text=is_text, is_pic=is_pic,
            is_video=is_video, is_music=is_music,
            user_name=user_name,
            min_words=min_words,
            stop_event=stop_event,
        )
        if not weibo_ids:
            logger.warning(
                f"未收集到任何微博:博主 {user_name}({user_id}) 在 "
                f"{start_date} ~ {end_date} 期间没有符合条件(原创长文)的微博"
                f"(可能停更/禁言或时间范围不符)。请确认时间范围后重试。\n"
                f"提示:如该时间段实际已爬取过,可取消勾选“跳过已爬取的文章”"
                f"强制重新处理。")
            return result
        if username and username != user_id:
            result["username"] = username

        # 2. 保存 ID 到 txt(新结构: 放入年份目录,如 2026年/)
        #    文件名沿用用户输入的原始日期
        data_dir = os.path.join(data_root, f"{result['username']}_{user_id}")
        start_year = start_date[:4]
        year_dir = os.path.join(data_dir, f"{start_year}年")
        os.makedirs(year_dir, exist_ok=True)
        txt_file = os.path.join(year_dir, f"{result['username']}_{user_id}_{start_date}_{end_date}.txt")
        crawler.save_weibo_ids_to_file(weibo_ids, txt_file)
        result["weibo_ids"] = weibo_ids
        result["txt_file"] = txt_file

        if skip_export:
            return result

        # 3. 导出(新结构: 博主目录/年份/月份/文件,重跑时覆盖同名文件)
        result["md_dir"] = data_dir
        stats_store = WeiboStatsStore(data_dir)
        ok_count, fail_count, skipped_count, ai_skipped_count = crawler.export_markdowns(
            weibo_ids, user_id, result["username"], data_dir,
            progress_callback=progress_callback, overwrite=True,
            month_subdirs=True,
            min_interval=min_interval, max_interval=max_interval,
            download_images=download_images, download_videos=download_videos,
            export_format=export_format, skip_existing=skip_existing,
            min_words=min_words,
            ai_enabled=ai_enabled, ai_classifier=ai_classifier,
            ai_config=ai_config, ai_usage_callback=ai_usage_callback,
            ai_rename=ai_rename,
            ai_root=ai_root or os.path.join(app_dir(), "AI分类"),
            ai_copy_media=ai_copy_media,
            stop_event=stop_event,
            stats_store=stats_store,
            force_rejudge=force_rejudge)
        result["exported"] = ok_count
        result["failed"] = fail_count
        result["skipped"] = skipped_count
        result["ai_skipped"] = ai_skipped_count
        result["ai_tokens"] = ai_tokens[0]
        result["stats_file"] = stats_store.path

    except RuntimeError as e:
        # 预期内的业务错误(如博主该时间段无微博、类名无法确定),只提示不打印堆栈
        logger.error(f"任务未完成: {e}")
        result["error"] = str(e)
    except Exception as e:
        logger.error(f"程序执行过程中出现错误: {e}", exc_info=True)
        result["error"] = str(e)
    finally:
        result["stopped"] = bool(stop_event is not None and stop_event.is_set())
        if not keep_browser_open:
            crawler.close(force_close=True)

    return result


# ---------------------------------------------------------------------------
# 筛选功能:对本地已爬取的数据文件按转评赞之和排序筛选
# ---------------------------------------------------------------------------

# 统计数字解析: 支持 "1.5万" / "3504.7万" / "123" / "1.2亿" 等微博常用格式
_COUNT_RE = re.compile(r"([\d.]+)\s*(万|亿)?")


def parse_count(text):
    """解析微博统计数字,返回 int;解析失败返回 0"""
    if not text:
        return 0
    m = _COUNT_RE.search(str(text).replace(",", "").strip())
    if not m:
        return 0
    num = float(m.group(1))
    unit = m.group(2) or ""
    if unit == "万":
        return int(num * 10000)
    if unit == "亿":
        return int(num * 100000000)
    return int(num)


class ArticleFilter:
    """本地文章筛选器

    扫描 DataPC/<博主>_<ID>/ 下各年份/月份文件夹中的 md/docx 文件,
    解析每篇的转发/评论/点赞数,按(可勾选的)数量之和排序,
    复制得分最高的前 N 篇到"筛选"文件夹,并按排名重命名。
    """

    STAT_KEYS = ("repost", "comment", "like")
    STAT_LABELS = {"repost": "转发", "comment": "评论", "like": "点赞"}
    # md 中的统计行: >转发数：331 / >评论数：225 / >点赞数：5084
    MD_STAT_RE = re.compile(r">(转发数|评论数|点赞数)[:：]\s*([\d.,万亿]+)")
    # 数据源: 键为 GUI 下拉框取值,值为 AI 分类目录后缀(位于"筛选"文件夹下)
    AI_SOURCE_LABELS = {
        "ai_high": "高质量",
        "ai_ad": "广告",
        "ai_suspicious": "可疑",
    }

    def __init__(self, data_root=None, filter_root="筛选", ai_root=None):
        self.data_root = data_root or os.path.join(app_dir(), "DataPC")
        self.filter_root = filter_root
        # AI 分类输出目录(与热度/字数筛选分开放;搜索时兼容旧位置"筛选"目录)
        self.ai_root = ai_root or os.path.join(app_dir(), "AI分类")

    # ---------- 扫描与解析 ----------

    def list_bloggers(self):
        """列出 DataPC 下已有的博主目录,返回 [(昵称, ID), ...](跳过空目录)"""
        result = []
        if not os.path.isdir(self.data_root):
            return result
        for name in sorted(os.listdir(self.data_root)):
            m = re.match(r"^(.+)_(\d+)$", name)
            p = os.path.join(self.data_root, name)
            if m and os.path.isdir(p):
                # 跳过空壳目录(如 get_username 失败创建的 unknown_ID)
                n_files = sum(len(fs) for _, _, fs in os.walk(p))
                if n_files > 0:
                    result.append((m.group(1), m.group(2)))
        return result

    def _find_user_dir(self, user_id):
        """在 DataPC 下查找指定 user_id 的数据目录

        优先选择包含"年份年"子目录的目录(真正的数据目录);
        若有多个匹配(如 unknown_ID 空壳),选内容最多的。
        """
        candidates = []
        for name in os.listdir(self.data_root):
            m = re.match(r"^(.+)_(\d+)$", name)
            if m and m.group(2) == user_id:
                p = os.path.join(self.data_root, name)
                if os.path.isdir(p):
                    candidates.append(p)
        if not candidates:
            return None
        # 按"年份年"子目录数量 + 文件总数 排序,取最像数据目录的
        def weight(p):
            year_dirs = sum(1 for x in os.listdir(p)
                            if os.path.isdir(os.path.join(p, x)) and x.endswith("年"))
            n_files = sum(len(fs) for _, _, fs in os.walk(p))
            return (year_dirs, n_files)
        return max(candidates, key=weight)

    def _find_ai_dir(self, user_name, source_key):
        """在 AI 分类目录(新: 程序目录/AI分类)或筛选目录(旧: 筛选/AI_*)下查找分类目录

        目录: AI_<博主名>_<类别>/<年份>年/<月份>/
        返回目录路径;未找到返回 None
        """
        label = self.AI_SOURCE_LABELS.get(source_key)
        if not label:
            return None
        target = f"AI_{user_name}_{label}"
        for root in (self.filter_root, self.ai_root):
            try:
                for name in os.listdir(root):
                    if name == target and os.path.isdir(os.path.join(root, name)):
                        return os.path.join(root, name)
            except Exception:
                continue
        return None

    def scan_files(self, user_id, start_date, end_date, source_format="md",
                   data_source="DataPC", user_name=None):
        """扫描指定博主、日期范围内、指定格式的文章文件

        data_source 为 "DataPC" 时扫描 DataPC/<博主>_<ID> 目录树;
        为 AI 数据源(ai_high/ai_ad/ai_suspicious)时扫描
        筛选/AI_<博主名>_<类别> 目录树。

        返回 [(file_path, file_date(datetime)), ...]
        """
        if data_source and data_source != "DataPC":
            user_dir = self._find_ai_dir(user_name or "", data_source)
            if not user_dir:
                logger.warning(
                    f"未找到AI分类目录: 筛选/AI_{user_name}_{self.AI_SOURCE_LABELS.get(data_source, '')}")
                return []
        else:
            user_dir = self._find_user_dir(user_id)
        if not user_dir:
            return []

        start = datetime.strptime(start_date, "%Y-%m-%d") if start_date else None
        end = datetime.strptime(end_date, "%Y-%m-%d") if end_date else None
        if start and end and start > end:
            logger.warning(f"扫描范围无效: 开始日期({start_date})晚于结束日期({end_date})")
            return []
        ext = ".docx" if source_format == "docx" else ".md"
        found = []
        for year_item in sorted(os.listdir(user_dir)):
            year_path = os.path.join(user_dir, year_item)
            if not (os.path.isdir(year_path) and year_item.endswith("年")):
                continue
            for month_item in sorted(os.listdir(year_path)):
                month_path = os.path.join(year_path, month_item)
                if not (os.path.isdir(month_path) and month_item.endswith("月")):
                    continue
                for f in os.listdir(month_path):
                    if not f.endswith(ext):
                        continue
                    # 兼容三种文件名:
                    #   新: <日期>_<AI标题>.ext (如 26-1-22_深度解析.md)
                    #   旧: <博主>_<日期>_<ID>.ext 与 <AI标题>_<日期>.ext
                    # 优先取"开头是日期"的匹配;否则取下划线日期形式的第一个匹配
                    fms = list(re.finditer(
                        r"(?:^|_)(\d{2,4})-(\d{1,2})-(\d{1,2})(?=_|\.(?:md|docx)$)", f))
                    if not fms:
                        continue
                    fm = fms[0]
                    year = int(fm.group(1))
                    if year < 100:
                        year += 2000
                    fdate = datetime(year, int(fm.group(2)), int(fm.group(3)))
                    if start and fdate < start:
                        continue
                    if end and fdate > end:
                        continue
                    found.append((os.path.join(month_path, f), fdate))
        if not found:
            try:
                n_all = sum(len(fs) for _, _, fs in os.walk(user_dir))
            except Exception:
                n_all = 0
            logger.warning(
                f"在 {user_dir} 中共 {n_all} 个文件,但未找到 "
                f"{start_date} ~ {end_date} 的 {ext} 文件"
                f"(若为AI重命名文件,请确认标题后带有 _年-月-日 日期)")
        return found

    # ---------- 统计解析 ----------

    def read_stats(self, file_path):
        """读取单个文件的转评赞与正文字数,返回 (repost, comment, like, word_count)"""
        repost = comment = like = 0
        word_count = 0
        try:
            if file_path.endswith(".md"):
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                for label, key in (("转发数", "repost"), ("评论数", "comment"),
                                   ("点赞数", "like"), ("正文字数", "word_count")):
                    m = re.search(rf">\s*{label}\s*[:：]\s*([\d.,万亿]+)", text)
                    if m:
                        val = parse_count(m.group(1))
                        if key == "repost":
                            repost = val
                        elif key == "comment":
                            comment = val
                        elif key == "like":
                            like = val
                        else:
                            word_count = val
            elif file_path.endswith(".docx"):
                from docx import Document
                doc = Document(file_path)
                for p in doc.paragraphs:
                    t = p.text.strip()
                    m = re.match(r"^(转发数|评论数|点赞数|正文字数)[:：]\s*([\d.,万亿]+)$", t)
                    if m:
                        val = parse_count(m.group(2))
                        if m.group(1) == "转发数":
                            repost = val
                        elif m.group(1) == "评论数":
                            comment = val
                        elif m.group(1) == "点赞数":
                            like = val
                        else:
                            word_count = val
        except Exception as e:
            logger.warning(f"读取统计失败 {file_path}: {e}")
        return repost, comment, like, word_count

    # ---------- 主流程 ----------

    def filter_top(self, user_name, user_id, start_date, end_date,
                   top_n=10, use_repost=True, use_comment=True, use_like=True,
                   source_format="md", move=False, by_word_count=False,
                   data_source="DataPC"):
        """按所选指标排序筛选前 N 篇,复制到"筛选"文件夹

        by_word_count=True 时按正文字数排序(忽略转评赞勾选);
        否则按转评赞之和排序。
        data_source 为 AI 数据源时,从"筛选/AI_<博主>_<类别>"目录中筛选。

        返回 dict: {"output_dir", "items": [...], "skipped": [...]}
        """
        ensure_logger()
        # 1. 扫描文件
        files = self.scan_files(user_id, start_date, end_date, source_format,
                                data_source=data_source, user_name=user_name)
        if not files:
            if data_source and data_source != "DataPC":
                logger.warning(
                    f"在 {self.filter_root} 的 AI_{user_name}_* 目录中未找到 "
                    f"{start_date} ~ {end_date} 的 {source_format} 文件")
            else:
                logger.warning(
                    f"在 {self.data_root} 中未找到博主 {user_name}({user_id}) "
                    f"{start_date} ~ {end_date} 的 {source_format} 文件")
            return {"output_dir": None, "items": [], "skipped": []}

        # 2. 解析统计并计算得分
        # 修复: 未勾选任何指标时默认按转评赞之和排序(此前会退化为按日期排序)
        if not by_word_count and not (use_repost or use_comment or use_like):
            use_repost = use_comment = use_like = True
            logger.info("未勾选任何指标,默认按 转发+评论+点赞 之和排序")
        records = []
        for fpath, fdate in files:
            repost, comment, like, word_count = self.read_stats(fpath)
            if by_word_count:
                score = word_count
            else:
                score = 0
                if use_repost:
                    score += repost
                if use_comment:
                    score += comment
                if use_like:
                    score += like
            records.append({
                "file": fpath, "date": fdate,
                "repost": repost, "comment": comment, "like": like,
                "word_count": word_count, "score": score,
                "by_word_count": by_word_count,
            })

        # 3. 排序取前 N(得分降序;同分按日期新->旧)
        records.sort(key=lambda r: (-r["score"], -r["date"].timestamp()))
        top = records[:top_n]

        # 4. 创建输出文件夹: 筛选/<博主名>_<ID>/<起>~<止>_热度TOP<N> 或 _字数TOP<N>
        #    文件夹名包含 数据源 与 勾选指标,避免不同数据源/不同指标组合的
        #    筛选结果混到同一个文件夹(如 DataPC 与 AI分类的"热度TOP10"分开)。
        #    格式: <起>~<止>_热度TOP<N>[_数据源][_指标]
        #    日期年份省略"20"(如 26-08-01),让文件夹名短一点;
        #    数据源为默认 DataPC 时省略,AI 数据源追加 _AI_类别;
        #    热度排序且未勾全三项时追加指标(如 _转发,_转发评论)。
        #    (先按博主分一层,避免多位博主的结果混在一起)
        os.makedirs(self.filter_root, exist_ok=True)
        metric_tag = "字数" if by_word_count else "热度"

        def _short_date(s):
            s = (s or "").strip()
            return s[2:] if len(s) >= 4 and s.startswith("20") else s

        date_tag = f"{_short_date(start_date)}~{_short_date(end_date)}"
        src_tag = {
            "DataPC": "",
            "ai_high": "AI_高质量",
            "ai_ad": "AI_广告",
            "ai_suspicious": "AI_可疑",
        }.get(data_source or "DataPC", "")
        if by_word_count:
            ind_tag = ""
        else:
            labels = []
            if use_repost:
                labels.append("转发")
            if use_comment:
                labels.append("评论")
            if use_like:
                labels.append("点赞")
            ind_tag = "" if len(labels) >= 3 else ("".join(labels) or "全部")
        name_parts = [f"{date_tag}_{metric_tag}TOP{len(top)}", src_tag, ind_tag]
        folder_name = "_".join(p for p in name_parts if p)
        out_dir = os.path.join(
            self.filter_root,
            f"{user_name}_{user_id}",
            folder_name)
        os.makedirs(out_dir, exist_ok=True)

        # 5. 复制/移动并重命名(序号前缀体现排名),同步复制图片/视频
        items = []
        for idx, rec in enumerate(top, 1):
            src = rec["file"]
            base = os.path.basename(src)
            name, ext = os.path.splitext(base)
            # 序号_原始名_得分.ext,如 01_卢诗翰_26-4-7_xxx_8888.md
            new_name = f"{idx:02d}_{name}_{rec['score']}{ext}"
            dst = os.path.join(out_dir, new_name)
            if move:
                shutil.move(src, dst)
            else:
                shutil.copy2(src, dst)
            # 同步复制同目录 images/ videos/ 下属于该微博的媒体文件
            self._copy_weibo_media(src, out_dir, move)
            rec["new_name"] = new_name
            items.append(rec)

        # 6. 生成统计说明文件
        self._write_summary(out_dir, items, use_repost, use_comment, use_like,
                            start_date, end_date, top_n, by_word_count)

        logger.info(f"筛选完成: 共扫描 {len(records)} 篇, 输出前 {len(top)} 篇到 {out_dir}")
        return {"output_dir": out_dir, "items": items, "skipped": []}

    @staticmethod
    def _copy_weibo_media(src_file, out_dir, move=False):
        """把与文章同目录 images/ videos/ 下属于该微博的媒体文件复制/移动到输出目录

        媒体文件名格式: <微博ID>_<序号>.<ext>(如 R9UIBEhFF_1.jpg)
        从文章文件名中提取微博ID(文件名形如 ..._<日期>_<微博ID>.md)
        """
        try:
            m = re.search(r"_([A-Za-z0-9]+)\.(?:md|docx)$", os.path.basename(src_file))
            if not m:
                return
            weibo_id = m.group(1)
            src_dir = os.path.dirname(src_file)
            for sub in ("images", "videos"):
                src_sub = os.path.join(src_dir, sub)
                if not os.path.isdir(src_sub):
                    continue
                dst_sub = os.path.join(out_dir, sub)
                os.makedirs(dst_sub, exist_ok=True)
                prefix = f"{weibo_id}_"
                for fname in os.listdir(src_sub):
                    if fname.startswith(prefix):
                        s = os.path.join(src_sub, fname)
                        d = os.path.join(dst_sub, fname)
                        if move:
                            shutil.move(s, d)
                        else:
                            shutil.copy2(s, d)
        except Exception as e:
            logger.warning(f"复制微博媒体失败 {src_file}: {e}")

    @staticmethod
    def _write_summary(out_dir, items, use_repost, use_comment, use_like,
                       start_date, end_date, top_n, by_word_count=False):
        """在输出文件夹中生成 筛选说明.txt,列出排名与各项数据"""
        if by_word_count:
            metric = "正文字数"
            header = "排名\t字数\t转发\t评论\t点赞\t文件名"
            row = lambda i, r: (
                f"{i}\t{r['word_count']}\t{r['repost']}\t{r['comment']}\t"
                f"{r['like']}\t{r['new_name']}")
        else:
            labels = []
            if use_repost:
                labels.append("转发")
            if use_comment:
                labels.append("评论")
            if use_like:
                labels.append("点赞")
            metric = "+".join(labels) if labels else "(未勾选任何指标)"
            header = "排名\t得分\t转发\t评论\t点赞\t文件名"
            row = lambda i, r: (
                f"{i}\t{r['score']}\t{r['repost']}\t{r['comment']}\t"
                f"{r['like']}\t{r['new_name']}")

        lines = [
            f"筛选范围: {start_date} ~ {end_date}",
            f"排序指标: {metric}",
            f"输出篇数: {min(top_n, len(items))}",
            "",
            header,
        ]
        for i, rec in enumerate(items, 1):
            lines.append(row(i, rec))
        try:
            with open(os.path.join(out_dir, "筛选说明.txt"),
                      "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
        except Exception as e:
            logger.warning(f"写入筛选说明失败: {e}")


if __name__ == "__main__":
    # 简单自测:查看默认类名配置
    ensure_logger()
    cm = ClassNameManager()
    logger.info(f"当前类名配置: {json.dumps(cm.classes, ensure_ascii=False, indent=2)}")
